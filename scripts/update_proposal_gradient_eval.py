from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Agentic_RAG_Project_Proposal_EN.docx"
GRADIENT_JSON = ROOT / "eval" / "results" / "agentic_gradient.json"

VARIANT_LABELS = {
    "naive": "E0 Retrieval",
    "e1": "E1 Routed/Tools",
    "e2": "E2 Graded",
    "full": "E3 Full Agentic",
}
VARIANTS = ["naive", "e1", "e2", "full"]


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell: Any, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table: Any, color: str = "D9E2F3") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def apply_table_style(table: Any, *, widths: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
                    set_cell_width(row.cells[idx], int(width * 1440))

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5 if row_idx else 8)
            if row_idx == 0:
                set_cell_shading(cell, "EAF2FF")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 78, 121)


def insert_paragraph_after(paragraph: Any, text: str = "", style: str | None = None) -> Any:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Any, rows: int, cols: int) -> Any:
    body = paragraph._parent
    table = body.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addnext(table._tbl)
    return table


def next_paragraph(paragraph: Any, text: str = "", style: str | None = None) -> Any:
    return insert_paragraph_after(paragraph, text, style)


def replace_text(paragraph: Any, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def score(value: float) -> str:
    return f"{value:.1f}"


def find_paragraph(doc: Any, startswith: str) -> Any:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Paragraph not found: {startswith}")


def update_variant_table(doc: Any, metrics: dict[str, dict[str, float]]) -> None:
    target = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "Variant":
            target = table
            break
    if target is None:
        raise ValueError("Variant table not found")

    updates = {
        1: (
            "Score 25.3; baseline for comparison",
            "Retrieval-only answers cannot access live business records and often fall back to generic policy context.",
        ),
        2: (
            "Score 50.6; action accuracy 100%",
            "Complete IDs call tools, incomplete IDs are rejected safely, and handoff intent routes to human support.",
        ),
        3: (
            "Score 70.6; evidence grading 100% on governed KB cases",
            "Safety and SOP answers include grade_docs before generation, improving evidence control.",
        ),
        4: (
            "Score 88.8; reflection observed in 90.9% of cases",
            "The full path adds post-generation reflection while preserving routing, tool use, and grading.",
        ),
    }
    for row_idx, (result, evidence) in updates.items():
        target.rows[row_idx].cells[3].text = result
        target.rows[row_idx].cells[4].text = evidence
    apply_table_style(target, widths=[1.05, 1.65, 1.25, 1.35, 2.2])


def main() -> None:
    payload = json.loads(GRADIENT_JSON.read_text(encoding="utf-8"))
    metrics = payload["metrics"]
    doc = Document(DOCX_PATH)

    p40 = find_paragraph(doc, "The evaluation uses four ablation variants.")
    replace_text(
        p40,
        "The evaluation uses four ablation variants. E0 is a naive retrieval-only baseline, "
        "E1 adds routing, deterministic business tools, human handoff, and governed retrieval, "
        "E2 adds evidence grading, and E3 adds post-generation reflection. The evaluation package "
        "contains a 30-case governed retrieval set, an 80-question agentic showcase set, and a staged "
        "gradient test that runs the same sampled questions across E0, E1, E2, and E3.",
    )

    p41 = find_paragraph(doc, "Expected conclusion:")
    replace_text(
        p41,
        "Observed conclusion: the latest staged gradient test produced a monotonic capability curve "
        "across the four variants: 25.3 -> 50.6 -> 70.6 -> 88.8. E1 separates the system from naive "
        "RAG through routing, tool calls, incomplete-ID handling, and handoff. E2 adds explicit evidence "
        "grading for governed knowledge-base cases. E3 adds reflection, making the final workflow more "
        "auditable and easier to demonstrate.",
    )

    update_variant_table(doc, metrics)

    cursor = p41
    cursor = next_paragraph(cursor, "Latest staged ablation results", "Heading 3")
    cursor = next_paragraph(
        cursor,
        "The staged gradient evaluation was run on a stratified 22-question sample covering order, "
        "logistics, after-sales, invoice, complaint, incomplete-ID, handoff, safety, SOP, requirements, "
        "and table-style questions. Each question was executed on all four variants, producing 88 "
        "end-to-end runs.",
        "Normal",
    )

    result_table = insert_table_after(cursor, rows=1 + len(VARIANTS), cols=7)
    headers = ["Stage", "Action", "Governed Context", "Evidence Grade", "Reflection", "No Error", "Score"]
    for idx, header in enumerate(headers):
        result_table.cell(0, idx).text = header
    for row_idx, variant in enumerate(VARIANTS, start=1):
        m = metrics[variant]
        values = [
            VARIANT_LABELS[variant],
            pct(m["action_accuracy"]),
            pct(m["governed_context_hit"]),
            pct(m["evidence_grade_rate"]),
            pct(m["reflection_rate"]),
            pct(m["no_error_rate"]),
            score(m["capability_score"]),
        ]
        for col_idx, value in enumerate(values):
            result_table.cell(row_idx, col_idx).text = value
    apply_table_style(result_table, widths=[1.3, 0.85, 1.25, 1.1, 0.95, 0.85, 0.75])

    cursor = result_table._element
    paragraph_after_table = doc.add_paragraph()
    cursor.addnext(paragraph_after_table._p)
    paragraph_after_table.style = "Normal"
    paragraph_after_table.add_run(
        "Case-level reporting is intentionally detailed. The generated report records each case's expected "
        "signal, observed signal, route, graph steps, tool status, latency, top evidence, and answer preview. "
        "This makes the evaluation useful for defense presentation as well as debugging: live-status questions "
        "show the jump from retrieval to tools, safety/SOP questions show doc_type-governed context selection, "
        "E2 exposes evidence grading, and E3 exposes reflection."
    )

    cursor_p = paragraph_after_table
    cursor_p = next_paragraph(cursor_p, "Representative case signals", "Heading 3")
    case_table = insert_table_after(cursor_p, rows=5, cols=4)
    case_headers = ["Case Type", "Question Example", "Observed Gradient", "Interpretation"]
    for idx, header in enumerate(case_headers):
        case_table.cell(0, idx).text = header
    examples = [
        (
            "Live order/tool lookup",
            "Order A1100 courier and tracking number",
            "E0: retrieve->generate; E1/E2: router->tool_call->generate; E3 adds reflect",
            "Agentic routing prevents a static policy answer from replacing live business-state lookup.",
        ),
        (
            "Safety boundary",
            "Approve after-sales review without proof",
            "E0: safety not frontloaded; E1: safety context first; E2: grade_docs; E3: reflect",
            "Safety documents are prioritized before ordinary policy answers when the request is risky.",
        ),
        (
            "Human escalation",
            "Angry user asks for a supervisor",
            "E0: retrieve->generate; E1/E2/E3: router->human_handoff",
            "The system can decide not to answer and instead trigger an operational handoff path.",
        ),
        (
            "Known optimization target",
            "Requirements/table cases",
            "Some table/requirements cases still miss governed context in top-3 evidence",
            "The test exposes remaining retrieval tuning work instead of hiding it behind answer fluency.",
        ),
    ]
    for row_idx, row in enumerate(examples, start=1):
        for col_idx, value in enumerate(row):
            case_table.cell(row_idx, col_idx).text = value
    apply_table_style(case_table, widths=[1.15, 1.55, 2.25, 2.05])

    doc.save(DOCX_PATH)
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
