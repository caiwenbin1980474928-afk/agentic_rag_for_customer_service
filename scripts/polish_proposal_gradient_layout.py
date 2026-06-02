from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Agentic_RAG_Project_Proposal_EN.docx"


def set_cell_margins(cell: Any, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def polish_tables(doc: Any) -> None:
    for table in doc.tables:
        header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        is_eval_table = (
            header.startswith("Variant |")
            or header.startswith("Stage |")
            or header.startswith("Case Type |")
        )
        if not is_eval_table:
            continue
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.05
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(8 if row_idx == 0 else 8.5)


def polish_paragraphs(doc: Any) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Observed conclusion:"):
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if text.startswith("Case-level reporting is intentionally detailed."):
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def main() -> None:
    doc = Document(DOCX_PATH)
    polish_tables(doc)
    polish_paragraphs(doc)
    doc.save(DOCX_PATH)
    print(f"Polished {DOCX_PATH}")


if __name__ == "__main__":
    main()
