from pathlib import Path
import importlib.util

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_SCRIPT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/scripts/build_speaker_script.py")
DETAIL_SCRIPT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/scripts/build_detailed_speaker_script.py")
OUT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/output")
DOCX = OUT / "RAG_中英双语展示讲稿_口播详细版.docx"
MD = OUT / "RAG_中英双语展示讲稿_口播详细版.md"


def load_module(path, name):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


base = load_module(BASE_SCRIPT, "base_script")
detail = load_module(DETAIL_SCRIPT, "detail_script")
slides = base.slides
details = detail.details


cn_replacements = [
    ("这一页还可以先帮听众建立一个直觉：", "这一页，我们先建立一个直觉："),
    ("开场时建议强调本次汇报的主线：", "开场时，我想先把本次汇报的主线说清楚："),
    ("这一页可以把“动机”讲得更业务化一些。", "这一页，我们把“动机”讲得更业务化一点。"),
    ("定义这一页建议按流程逐格讲。", "定义这一页，我们按流程一步一步看。"),
    ("讲五类问题时，可以把它们和真实风险对应起来。", "讲这五类问题时，我们把它们和真实风险对应起来看。"),
    ("这里可以用“脑内知识”和“外部资料库”做类比。", "这里我用“脑内知识”和“外部资料库”做一个类比。"),
    ("核心链路这一页可以先讲一个重要判断：", "核心链路这一页，我想先讲一个重要判断："),
    ("端到端链路可以按“数据变成证据，证据变成答案”来讲。", "端到端链路，我们可以理解成“数据变成证据，证据再变成答案”。"),
    ("Chunking 可以结合一个例子讲：", "Chunking 这个概念，我们用一个例子来看："),
    ("技术范式部分可以告诉听众：", "技术范式这一部分，我想先说明："),
    ("演进路线可以讲成能力逐步叠加。", "这条演进路线，本质上是能力一步一步叠加。"),
    ("范式地图可以用来避免概念混淆。", "这张范式地图，主要是帮助我们避免概念混淆。"),
    ("这个案例可以讲得更像真实客服流程。", "这个案例，我们就把它当成一个真实客服流程来看。"),
    ("应用、挑战和评估这一部分可以作为落地总结。", "应用、挑战和评估这一部分，是对落地问题的总结。"),
    ("典型场景这一页可以按“为什么适合”和“风险边界”两条线讲。", "典型场景这一页，我们按两条线来看：为什么适合，以及风险边界在哪里。"),
    ("优势这一页可以总结成四个关键词：", "优势这一页，我把它总结成四个关键词："),
    ("挑战这一页适合强调“不要把问题都归因给模型”。", "挑战这一页，我想强调一点：不要把问题都归因给模型。"),
    ("评估这一页要讲清楚“答案对了”并不等于系统没问题。", "评估这一页，我们一定要区分一件事：“答案看起来对了”，并不等于系统没有问题。"),
    ("这一页可以帮助听众避免技术路线混用。", "这一页，我们把 RAG 和其他技术的分工讲清楚，避免把技术路线混在一起。"),
    ("最后一页可以把未来趋势和前面的主线连接起来。", "最后一页，我们把未来趋势和前面的主线连起来看。"),
    ("Agent Loop 可以讲成观察、判断、行动的闭环。", "Agent Loop 本质上是一个观察、判断、行动的闭环。"),
    ("建议讲一个排查顺序：", "这里我会按一个排查顺序来讲："),
    ("可以举一个简单例子：", "举一个简单例子："),
    ("也可以提醒：", "这里也要提醒大家："),
    ("可以说，", "也就是说，"),
    ("可以用", "我们用"),
    ("建议", "我建议"),
    ("提醒听众", "提醒大家"),
    ("这页的结论是：", "所以这一页的结论是："),
    ("这页可以强调", "这里我想强调"),
    ("这里要提醒：", "这里要提醒大家："),
    ("不要讲成", "不要把它讲成"),
]


en_replacements = [
    ("At this point, it is useful to give the audience an intuition:", "At this point, I want to give you a simple intuition:"),
    ("For the opening, emphasize the storyline of the talk:", "So the storyline for this talk is:"),
    ("This slide can be framed in business terms.", "On this slide, I want to frame the motivation in business terms."),
    ("For this definition slide, walk through the flow step by step.", "For this definition slide, let us walk through the flow step by step."),
    ("When explaining the five problems, connect each one to a real risk.", "When we look at these five problems, we should connect each one to a real risk."),
    ("A useful analogy is", "A useful way to think about this is"),
    ("For the core pipeline slide, start with an important observation:", "For the core pipeline, I want to start with an important observation:"),
    ("The end-to-end pipeline can be explained as", "The end-to-end pipeline is basically"),
    ("Chunking can be explained with a policy example.", "Let us make chunking concrete with a policy example."),
    ("In the technical paradigms section, tell the audience that", "In the technical paradigms section, the first point is that"),
    ("The evolution path can be explained as accumulating capabilities.", "This evolution path is really a story of accumulating capabilities."),
    ("The paradigm map helps avoid conceptual confusion.", "This paradigm map helps us avoid conceptual confusion."),
    ("This case can be presented like a real customer-service workflow.", "Let us treat this case like a real customer-service workflow."),
    ("The applications, challenges, and evaluation section can serve as the implementation summary.", "The applications, challenges, and evaluation section is the implementation summary."),
    ("For the use-case slide, explain each scenario through two lines:", "For the use-case slide, I would read each scenario through two lines:"),
    ("The advantages can be summarized with four keywords:", "I would summarize the advantages with four keywords:"),
    ("This challenge slide should emphasize that", "On the challenge slide, the key message is that"),
    ("This evaluation slide should make clear that", "On the evaluation slide, we need to make clear that"),
    ("This slide helps the audience avoid mixing up technical routes.", "This slide helps us avoid mixing up technical routes."),
    ("The final slide should connect future trends back to the main storyline.", "On the final slide, we connect future trends back to the main storyline."),
    ("This tool landscape slide should not be presented as a list of products.", "For the tool landscape, I do not want to present it as a list of products."),
    ("It should be presented as layer-based selection.", "Instead, I want to present it as layer-based selection."),
    ("A simple example is", "For example,"),
    ("For example, an employee asking,", "For example, if an employee asks,"),
    ("It is also important to remind the audience that", "It is also worth emphasizing that"),
    ("This is also a good place to remind the audience that", "This is also where I would emphasize that"),
    ("The key message is that", "The key message here is that"),
    ("This helps the audience see", "This helps us see"),
    ("This is also where I would emphasize that", "Here I want to emphasize that"),
    ("For the use-case slide, I would read each scenario through two lines:", "For these use cases, we can look at each scenario through two lines:"),
    ("I would summarize the advantages with four keywords:", "I summarize the advantages with four keywords:"),
    ("A strong closing sentence is:", "So I will close with this point:"),
]


def spoken_cn(text):
    out = text
    for old, new in cn_replacements:
        out = out.replace(old, new)
    out = out.replace("。RAG", "。RAG")
    return out


def spoken_en(text):
    out = text
    for old, new in en_replacements:
        out = out.replace(old, new)
    return out


def slide_paragraphs(slide, lang):
    parts = [slide[lang]] + details[slide["no"]][lang]
    transform = spoken_cn if lang == "cn" else spoken_en
    return [transform(p) for p in parts]


def set_font(run, size=None, bold=False, color=None, east_asia="STHeiti"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_para(doc, text, size=10.5, color=None, bold=False, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.24
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    color = "2E74B5" if level == 1 else "1F4D78"
    size = 15 if level == 1 else 12
    for run in p.runs:
        set_font(run, size=size, bold=True, color=color)
    return p


def build_md():
    lines = [
        "# RAG 中英双语展示讲稿（口播详细版）",
        "",
        "说明：本稿按英文版 PPT 的 33 页逐页对应。语言改为可直接发言的口播风格，去掉“建议讲、可以讲”一类备课提示语。",
        "",
    ]
    for s in slides:
        lines.extend([
            f"## Slide {s['no']} - {s['title']}",
            "",
            "### 中文口播稿",
            "",
            *slide_paragraphs(s, "cn"),
            "",
            "### English Spoken Script",
            "",
            *slide_paragraphs(s, "en"),
            "",
        ])
    return "\n\n".join(lines)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.24

    header = section.header.paragraphs[0]
    header.text = "RAG bilingual spoken speaker script"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=8.5, color="666666")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("RAG 中英双语展示讲稿（口播详细版）")
    set_font(r, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("逐页对应英文版 PPT，共 33 页。语言改为可直接发言的口播稿，每页中文与英文一一对应。")
    set_font(r, size=10.5, color="666666")

    meta = add_para(doc, "对应文件：RAG_Technical_Review_EN.pptx；处理原则：页序和主题不变，表达改成演讲现场可直接说出口的语言。", size=9.5, color="666666", after=10)
    shade_paragraph(meta, "F4F6F9")

    for s in slides:
        add_heading(doc, f"Slide {s['no']}  {s['title']}", 1)
        add_heading(doc, "中文口播稿", 2)
        for para in slide_paragraphs(s, "cn"):
            add_para(doc, para)
        add_heading(doc, "English Spoken Script", 2)
        for para in slide_paragraphs(s, "en"):
            add_para(doc, para)

    doc.save(DOCX)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    MD.write_text(build_md(), encoding="utf-8")
    build_docx()
    print(DOCX)
    print(MD)


if __name__ == "__main__":
    main()
