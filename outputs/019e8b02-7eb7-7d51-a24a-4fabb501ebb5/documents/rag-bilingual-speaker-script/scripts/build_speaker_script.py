from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/output")
DOCX = OUT / "RAG_中英双语展示讲稿.docx"
MD = OUT / "RAG_中英双语展示讲稿.md"


slides = [
    {
        "no": "01",
        "title": "Retrieval-Augmented Generation (RAG) Technical Review / 检索增强生成（RAG）技术综述",
        "cn": "这一页是整场展示的开场。我们先把主题定在检索增强生成，也就是 RAG。RAG 的核心并不是把所有知识都塞进大模型参数里，而是在模型回答前，为它提供来自外部系统的、可以追溯的证据。今天的内容会从概念与动机开始，逐步展开核心链路、检索技术、工程实现、技术范式、Agentic RAG，以及最后的评估和未来趋势。大家可以把 RAG 理解成一套把 LLM 和外部证据连接起来的工程框架：模型负责理解和生成，检索系统负责找到可靠依据。",
        "en": "This opening slide sets the scope for the whole presentation: Retrieval-Augmented Generation, or RAG. The key idea is not to make the model memorize everything in its parameters. Instead, before the model answers, we provide it with external, traceable evidence. The talk will move from concepts and motivation to the core pipeline, retrieval methods, engineering implementation, RAG paradigms, Agentic RAG, evaluation, and future trends. In simple terms, RAG is an engineering framework that connects an LLM with external evidence: the model handles understanding and generation, while the retrieval system finds reliable supporting information.",
    },
    {
        "no": "02",
        "title": "Concepts and Motivation / 概念与动机",
        "cn": "这一页进入第一部分：概念与动机。这里要回答三个问题：第一，RAG 到底是什么；第二，为什么只依赖模型参数记忆是不够的；第三，RAG 为什么在企业 AI 场景里特别重要。传统大模型虽然具备很强的语言理解、总结和推理能力，但它的知识主要来自训练阶段。企业真实场景里，知识会持续变化，很多信息也不在公开训练数据里，所以我们需要一种机制，让模型在推理阶段临时查证、引用和使用外部知识。",
        "en": "This slide introduces the first section: concepts and motivation. We need to answer three questions. First, what exactly is RAG? Second, why is model parameter memory alone not enough? Third, why is RAG especially important in enterprise AI? Large language models are strong at understanding, summarizing, and reasoning, but most of their knowledge comes from training. In real enterprise settings, knowledge changes continuously, and much of it is private. That is why we need a mechanism that lets the model look up, cite, and use external knowledge during inference.",
    },
    {
        "no": "03",
        "title": "Definition / RAG 的定义",
        "cn": "RAG 可以概括为：让大模型在生成答案之前先查资料，再基于证据回答。它不是一个单独的模型，而是由用户问题、外部知识检索、片段选择、Prompt 构造、LLM 生成和引用返回组成的一套流程。这里的关键是“证据在推理时进入上下文”。比如用户提出一个问题，系统先从文档、数据库或 API 中检索相关内容，筛选 Top-K 片段，再把这些片段和指令一起放入 Prompt，让模型基于这些资料作答，并返回来源、chunk 或链接。这样做的目标不是让模型记住所有企业知识，而是在需要时把证据交给模型。",
        "en": "RAG can be summarized as follows: the model looks up information before generating an answer, and then answers based on that evidence. It is not a single model. It is a workflow that includes the user question, external knowledge retrieval, chunk selection, prompt construction, LLM generation, and citation return. The key point is that evidence enters the model context at inference time. For example, when a user asks a question, the system retrieves relevant content from documents, databases, or APIs, selects Top-K chunks, places those chunks into the prompt with instructions, and asks the model to answer with sources, chunk IDs, or links. The goal is not to make the model memorize enterprise knowledge, but to supply evidence when needed.",
    },
    {
        "no": "04",
        "title": "Why RAG / 为什么需要 RAG",
        "cn": "这一页解释 RAG 解决的五类核心问题。第一是知识更新滞后，大模型训练后参数基本固定，新政策、新产品和实时订单状态很难及时进入模型。第二是幻觉风险，缺少证据时模型可能生成听起来合理但事实错误的答案。第三是私有知识缺失，合同、手册、代码和会议纪要通常不在公开训练数据中。第四是长上下文成本，如果把所有资料都塞进上下文，会带来成本、延迟和注意力噪声。第五是依据不透明，没有来源、段落和时间戳，就很难审计和追责。因此，RAG 的价值是把“回答”变成“基于证据的回答”。",
        "en": "This slide explains five core problems that RAG addresses. The first is stale knowledge: after training, model parameters are mostly fixed, so new policies, new products, and real-time order status are hard to include. The second is hallucination risk: without evidence, the model may produce answers that sound plausible but are factually wrong. The third is the private knowledge gap: contracts, manuals, code, and meeting notes are usually not in public training data. The fourth is long-context cost: putting everything into context increases cost, latency, and attention noise. The fifth is opaque evidence: without source, paragraph, and timestamp, answers are hard to audit. RAG turns an answer into an evidence-grounded answer.",
    },
    {
        "no": "05",
        "title": "Knowledge Types / 参数知识与非参数知识",
        "cn": "这里要区分两类知识。参数知识是模型训练后存储在参数里的知识，它通用、压缩、能支持语言理解和推理，但很难更新，也不能直接追溯到具体来源。非参数知识则是外部文档、数据库、搜索引擎、知识图谱等，它可更新、可审计、可控制，也可以按需检索。RAG 的本质就是把两者组合起来：模型用参数知识理解问题、组织语言和推理；系统用非参数知识提供最新、私有、可引用的事实依据。这个分工非常关键，因为它让知识更新从模型训练问题变成工程系统问题。",
        "en": "Here we distinguish two types of knowledge. Parametric knowledge is stored in the model parameters after training. It is general, compressed, and useful for language understanding and reasoning, but it is hard to update and not directly traceable. Non-parametric knowledge refers to external documents, databases, search engines, knowledge graphs, and other sources. It is updateable, auditable, controllable, and retrievable on demand. RAG combines the two: the model uses parametric knowledge to understand and reason, while the system uses non-parametric knowledge to provide current, private, and citable evidence. This division turns knowledge updating from a model-training problem into an engineering problem.",
    },
    {
        "no": "06",
        "title": "Core Pipeline / 核心链路",
        "cn": "第二部分开始讲 RAG 的核心链路。经典 RAG 可以拆成两条流水线：一条是离线索引，把知识整理成可检索的索引；另一条是在线问答，在用户提问时检索证据并生成答案。这里要强调，RAG 不是最后加一个 Prompt 就能解决的问题。离线索引、在线检索、上下文组织、答案生成和引用返回，每一步都会影响最终质量。如果前面解析错了、切分错了、召回不足，即使最后的模型很强，也很难给出可靠答案。",
        "en": "The second section begins with the core RAG pipeline. Classic RAG can be divided into two pipelines. The first is offline indexing, which turns knowledge into a retrievable index. The second is online question answering, where the system retrieves evidence and generates an answer when the user asks a question. The important point is that RAG is not solved by adding one prompt at the end. Offline indexing, online retrieval, context organization, answer generation, and citation return all affect quality. If parsing, chunking, or recall is poor, even a strong model will struggle to produce a reliable answer.",
    },
    {
        "no": "07",
        "title": "End-to-End Pipeline / 端到端链路",
        "cn": "这一页把生产级 RAG 的端到端流程串起来：采集、清洗、切分、向量化、索引、检索、生成。离线阶段决定知识能不能被找到，比如文档有没有被正确解析、噪声有没有清掉、chunk 是否合理、向量是否准确。在线阶段决定证据能不能被正确使用，比如检索 Top-K 是否足够、上下文是否去重排序、引用是否准确、模型是否忠于证据。很多 RAG 问题并不是生成模型本身造成的，而是链路中某个环节质量下降后传导到答案。",
        "en": "This slide connects the production RAG flow end to end: ingestion, cleaning, chunking, embedding, indexing, retrieval, and generation. The offline phase determines whether knowledge can be found: whether documents are parsed correctly, noise is removed, chunks are well formed, and vectors are meaningful. The online phase determines whether evidence can be used correctly: whether Top-K retrieval is sufficient, context is deduplicated and ordered, citations are accurate, and the model remains faithful to evidence. Many RAG failures are not caused by the generation model alone; they come from quality loss in earlier pipeline stages.",
    },
    {
        "no": "08",
        "title": "Offline Indexing / 离线索引",
        "cn": "离线索引更像数据工程，而不是 Prompt 工程。第一步是文档采集，来源可能包括 PDF、网页、数据库、API 和代码仓库，同时要处理权限过滤。第二步是文档清洗，包括去页眉页脚、去重复、表格解析和编码修复。第三步是切分，把长文档变成多个 chunk。第四步是向量化，用 Embedding 模型把文本片段转成向量。最后是向量存储，把向量和元数据写入 FAISS、Milvus、Chroma、Weaviate 或 Elasticsearch 等系统。只有离线证据库质量足够好，在线检索才有基础。",
        "en": "Offline indexing is closer to data engineering than prompt engineering. The first step is document ingestion, from sources such as PDFs, web pages, databases, APIs, and code repositories, while applying permission filters. The second step is document cleaning: removing headers and footers, deduplicating content, parsing tables, and fixing encoding issues. The third step is chunking, which turns long documents into retrievable chunks. The fourth step is embedding, where an embedding model converts text chunks into vectors. The final step is vector storage, writing vectors and metadata into systems such as FAISS, Milvus, Chroma, Weaviate, or Elasticsearch. High-quality offline evidence is the foundation for online retrieval.",
    },
    {
        "no": "09",
        "title": "Chunking / 文档切分",
        "cn": "Chunking 经常被低估，但它直接决定检索颗粒度和上下文可读性。固定长度切分简单稳定，但可能切断语义。滑动窗口通过重叠缓解边界丢失，但会增加存储和重复召回。按标题切分可以保留文档层级，但依赖原始文档结构。递归切分会按段落、句子、字符逐级处理，比较平衡但需要调参。语义切分更接近真实内容边界，但成本更高。对于表格和代码，还需要保留行列、函数、类等结构。切分的目标不是平均分块，而是让每个片段既完整又不过度噪声化。",
        "en": "Chunking is often underestimated, but it directly defines retrieval granularity and context readability. Fixed-length splitting is simple and stable, but it may cut through meaning. Sliding windows reduce boundary loss through overlap, but they increase storage and repeated recall. Heading-based splitting preserves document hierarchy, but it depends on source structure quality. Recursive splitting works through paragraphs, sentences, and characters, offering balance but requiring parameter tuning. Semantic splitting better follows content boundaries, but it is more expensive. Tables and code need special handling to preserve rows, columns, functions, and classes. The goal is not equal-sized chunks; it is chunks that are complete enough without adding too much noise.",
    },
    {
        "no": "10",
        "title": "Embedding / 向量化",
        "cn": "Embedding 决定语义检索能否理解领域语言。商业模型例如 OpenAI text-embedding 和 Cohere Embed，通常稳定、易用、效果强，但依赖 API。开源模型如 BGE、E5、GTE、Jina Embeddings，适合私有化部署和成本敏感场景。多模态模型如 CLIP、SigLIP，可以支持图片、截图、表格等内容检索。代码 Embedding 则适合代码库问答、API 文档和开发助手。选择 Embedding 时要关注是否理解缩写、同义词、跨语言表达和行业术语；如果这一步失败，正确文档可能一开始就召回不到。",
        "en": "Embedding quality determines whether semantic retrieval understands domain language. Commercial models such as OpenAI text-embedding and Cohere Embed are usually stable, easy to use, and high-performing, but they depend on APIs. Open-source models such as BGE, E5, GTE, and Jina Embeddings are useful for private deployment and cost-sensitive scenarios. Multimodal models such as CLIP and SigLIP support retrieval over images, screenshots, and tables. Code embeddings are useful for codebase QA, API documentation, and developer assistants. When choosing an embedding model, we need to check whether it understands abbreviations, synonyms, multilingual meaning, and industry terms. If this step fails, the right document may never be recalled.",
    },
    {
        "no": "11",
        "title": "Vector Database / 向量数据库",
        "cn": "当数据规模变大时，向量数据库要在准确率、延迟和成本之间做平衡。Flat 是逐个计算精确距离，准确但慢，适合小规模或评测。HNSW 是分层图结构搜索，查询快、召回高，但内存占用较大。IVF 是先聚类，再在相关簇中搜索，适合大规模并且可以调节召回和速度。PQ 或量化通过压缩向量降低存储成本，但可能损失精度。DiskANN 面向磁盘的大规模 ANN，适合超大规模低成本存储。选择索引时，不能只看速度，也要看召回、过滤能力、更新方式和部署成本。",
        "en": "As data grows, vector databases must balance accuracy, latency, and cost. Flat search computes exact distance against every vector. It is accurate but slow, so it is useful for small datasets or evaluation. HNSW uses hierarchical graph-based search; it is fast and has high recall, but it consumes more memory. IVF clusters vectors first and searches relevant clusters, making it suitable for large-scale data with tunable recall and speed. PQ or quantization compresses vectors to reduce storage cost, at the risk of accuracy loss. DiskANN is designed for large-scale disk-based ANN search. When choosing an index, we should consider not only speed, but also recall, filtering ability, update patterns, and deployment cost.",
    },
    {
        "no": "12",
        "title": "Retrieval / 检索方法",
        "cn": "这一页对比 BM25、Dense Retrieval 和 Hybrid Search。BM25 是稀疏检索，基于词项、倒排索引和词频，擅长精确术语、编号、代码和产品型号，但对没有词面重合的语义匹配较弱。Dense Retrieval 使用高维语义向量，擅长同义改写和自然语言问题，但可能忽略数字、罕见实体和硬性关键词。生产系统常用混合检索，把 BM25 的精确匹配和 Dense 的语义召回结合起来，再通过过滤、重排序和阈值控制结果质量。核心原则是，不同检索方式覆盖的是不同失败模式。",
        "en": "This slide compares BM25, dense retrieval, and hybrid search. BM25 is sparse retrieval based on terms, inverted indexes, and term frequency. It is strong at exact terms, IDs, code, and product models, but weaker when semantically relevant documents do not share words with the query. Dense retrieval uses high-dimensional semantic vectors. It is strong at synonyms, paraphrases, and natural-language questions, but it may miss numbers, rare entities, and hard lexical constraints. Production systems often combine both through hybrid search, then control quality with filters, reranking, and thresholds. The core principle is that different retrieval methods cover different failure modes.",
    },
    {
        "no": "13",
        "title": "Rerank and Prompt / 重排序与 Prompt 组织",
        "cn": "初召回的目标是尽量不漏，重排序和 Prompt 组织的目标是让证据可用。常见做法是 Retriever 先召回 Top-50，Reranker 再对 query-document pair 重新打分，保留最强的 Top-5 放入上下文。接着要做上下文拼接：去重、排序、压缩，并保留来源、段落和时间戳等元数据。Prompt 里要明确约束模型只能依据资料回答，证据不足时要说明无法确定。对于高风险问题，还要有拒答、澄清或转人工策略。成本上，则通过 Top-K、压缩、缓存和分层检索控制上下文长度与延迟。",
        "en": "Initial recall aims to avoid missing relevant evidence; reranking and prompt assembly make that evidence usable. A common pattern is that the retriever recalls Top-50 candidates, and the reranker scores query-document pairs again, keeping the strongest Top-5 for context. Then the system assembles context by deduplicating, ordering, compressing, and preserving metadata such as source, section, and timestamp. The prompt should clearly instruct the model to answer only from the supplied evidence and to say when the evidence is insufficient. For high-risk questions, the system also needs refusal, clarification, or human handoff. Cost is controlled through Top-K tuning, compression, caching, and hierarchical retrieval.",
    },
    {
        "no": "14",
        "title": "Engineering Implementation / 工程实现",
        "cn": "第三部分进入工程实现。一个可用的 RAG 系统不是单点能力，而是数据、索引、检索、生成、应用和监控层共同工作的系统。这里的重点是把 RAG 当成生产工程，而不是 Demo。Demo 只需要能回答几个问题，生产系统则要考虑数据同步、权限、增量更新、检索质量、引用约束、延迟成本、用户反馈和回归评估。后面两页会从工程分层和工具定位两个角度说明：每一层都有自己的职责，也有不同的工具和质量指标。",
        "en": "The third section moves into engineering implementation. A usable RAG system is not a single capability; it is a system where data, indexing, retrieval, generation, application, and monitoring layers work together. The key is to treat RAG as production engineering, not just a demo. A demo only needs to answer a few questions, but a production system must handle data synchronization, permissions, incremental updates, retrieval quality, citation control, latency, cost, user feedback, and regression evaluation. The next two slides explain this from two angles: system layers and tool roles. Each layer has its own responsibility, tools, and quality metrics.",
    },
    {
        "no": "15",
        "title": "Architecture / 系统架构",
        "cn": "一个实用 RAG 系统通常可以分成六层。数据准备层负责连接数据源、解析文档、清洗内容和提取元数据。索引构建层负责切分、Embedding、写入向量库和增量更新。检索层负责查询改写、BM25、向量或混合检索、过滤、重排序和压缩。生成层负责 Prompt、LLM 调用、引用控制和结构化输出。应用层负责会话、权限、UI、API、流式输出和人工兜底。评估监控层负责离线评测、线上监控、反馈、成本和延迟分析。排查问题时，也应该按这六层定位，而不是直接归因给模型。",
        "en": "A practical RAG system usually has six layers. The data preparation layer connects data sources, parses documents, cleans content, and extracts metadata. The index construction layer handles chunking, embedding, vector-store writes, and incremental updates. The retrieval layer handles query rewriting, BM25, vector or hybrid search, filters, reranking, and compression. The generation layer manages prompts, LLM calls, citation control, and structured output. The application layer handles conversation, permissions, UI, APIs, streaming, and human handoff. The evaluation and monitoring layer handles offline evaluation, online monitoring, feedback, cost, and latency analysis. When debugging, we should locate the failure by layer instead of blaming the model immediately.",
    },
    {
        "no": "16",
        "title": "Tool Landscape / 工具生态",
        "cn": "工具选型要按层理解，而不是把所有产品混成一类。LangChain 更偏编排、检索、工具调用和 Agent 应用搭建。LlamaIndex 强调私有数据连接、索引、查询引擎和评估。Haystack 适合模块化 QA Pipeline。FAISS 和 Milvus 解决向量索引与向量数据库问题。Elastic 和 Weaviate 覆盖全文、混合搜索和对象存储。BGE 和 E5 是 Embedding 模型，OpenAI 和 Qwen 等生成模型负责基于上下文输出答案。RAGAS、LangSmith、DeepEval 则用于评估、追踪、反馈和回归测试。选型前要先明确自己缺的是哪一层能力。",
        "en": "Tool selection should be understood by layer, not by treating all products as one category. LangChain focuses on orchestration, retrieval, tool use, and Agent application building blocks. LlamaIndex emphasizes private data connectors, indexing, query engines, and evaluation. Haystack is useful for modular QA pipelines. FAISS and Milvus address vector indexing and vector database infrastructure. Elastic and Weaviate cover full-text search, hybrid search, and object storage. BGE and E5 are embedding models, while generators such as OpenAI and Qwen produce answers from context. RAGAS, LangSmith, and DeepEval support evaluation, tracing, feedback, and regression testing. Before choosing tools, we should identify which layer is missing.",
    },
    {
        "no": "17",
        "title": "Technical Paradigms / 技术范式",
        "cn": "第四部分讲技术范式。RAG 正在从固定的“检索然后生成”链路，演化为更可组合、可纠错、可由 Agent 动态调度的知识架构。早期 RAG 更像一条线性流程，问题来了就检索，检索结果放进 Prompt，再让模型回答。后来的 Advanced RAG、Graph RAG、Modular RAG 和 Agentic RAG，则分别解决召回质量、关系理解、模块替换和动态决策等问题。理解这些范式不是为了追概念，而是为了知道在什么问题下该引入什么能力。",
        "en": "The fourth section discusses technical paradigms. RAG is evolving from a fixed retrieve-then-generate chain into a more composable, corrective, and Agent-orchestrated knowledge architecture. Early RAG looked like a linear flow: retrieve when a question arrives, put the retrieved content into the prompt, and ask the model to answer. Later paradigms such as Advanced RAG, Graph RAG, Modular RAG, and Agentic RAG address different problems: recall quality, relational understanding, module replacement, and dynamic decision-making. Understanding these paradigms is not about chasing terminology; it is about knowing which capability to introduce for which problem.",
    },
    {
        "no": "18",
        "title": "Evolution / RAG 演进",
        "cn": "RAG 的演进可以看作是在不断修复上一代系统的薄弱点。传统 IR 使用 TF-IDF、BM25 和倒排索引，速度快、可解释，但语义能力弱。Dense Retrieval 用神经网络向量支持语义相似。Classic RAG 把检索文档和 LLM 生成结合起来，降低幻觉并接入私有知识。Advanced RAG 加入查询改写、Rerank、压缩和过滤，提高上下文质量。Graph RAG 通过实体关系图和社区摘要补足全局关系分析。Agentic RAG 则让 Agent 规划、调用工具、重试和反思。能力越强，工程复杂度也越高。",
        "en": "The evolution of RAG can be seen as repeatedly fixing weaknesses of the previous generation. Traditional information retrieval uses TF-IDF, BM25, and inverted indexes. It is fast and explainable, but weak at semantics. Dense retrieval uses neural vectors to support semantic similarity. Classic RAG combines document retrieval with LLM generation, reducing hallucination and enabling private knowledge. Advanced RAG adds query rewriting, reranking, compression, and filtering to improve context quality. Graph RAG uses entity graphs and community summaries to support global relationship analysis. Agentic RAG lets Agents plan, call tools, retry, and reflect. As capability increases, engineering complexity also increases.",
    },
    {
        "no": "19",
        "title": "Paradigm Map / 范式地图",
        "cn": "这一页强调，不同 RAG 范式不是互斥选项，而是可组合能力。Naive RAG 固定执行检索后生成，简单低成本，但检索失败后恢复能力弱。Advanced RAG 加入改写、Rerank、压缩和过滤。Modular RAG 把流程拆成可替换模块。Corrective RAG 检查并纠正低质量召回。Self-RAG 让模型判断是否需要检索、答案是否充分。Adaptive RAG 根据问题复杂度选择策略。Graph RAG 组织实体关系，Agentic RAG 负责动态规划、工具调用和反思。生产系统通常会把这些能力组合起来，而不是只选一个标签。",
        "en": "This slide emphasizes that RAG paradigms are not mutually exclusive; they are composable capabilities. Naive RAG uses a fixed retrieve-then-generate flow. It is simple and low-cost, but weak at recovery when retrieval fails. Advanced RAG adds rewriting, reranking, compression, and filtering. Modular RAG splits the flow into replaceable modules. Corrective RAG checks and fixes poor recall. Self-RAG lets the model judge whether retrieval is needed and whether the answer is sufficient. Adaptive RAG chooses a strategy based on question complexity. Graph RAG organizes entity relationships, while Agentic RAG handles dynamic planning, tool use, and reflection. Production systems often combine these capabilities instead of choosing one label.",
    },
    {
        "no": "20",
        "title": "Graph and Modular RAG / Graph RAG 与 Modular RAG",
        "cn": "Graph RAG 和 Modular RAG 解决的是普通向量相似度难以覆盖的问题。Graph RAG 关注关系和全局主题，通过实体图、关系图和社区摘要，支持多文档关系推理、全局分析和主题总结。Modular RAG 关注流程可组合，把查询改写、检索器、过滤器、重排序、压缩器、生成器和评估器拆成可替换模块。代价是工程复杂度上升：图谱构建、增量更新、模块接口和评估都会变复杂。实际系统中，Advanced RAG 提升召回质量，Graph RAG 提升关系理解，Agentic RAG 决定何时调用哪条链路。",
        "en": "Graph RAG and Modular RAG solve problems that ordinary vector similarity does not cover well. Graph RAG focuses on relationships and global themes. Through entity graphs, relation graphs, and community summaries, it supports multi-document relationship reasoning, global analysis, and thematic summarization. Modular RAG focuses on composable workflows by separating query rewriting, retrievers, filters, rerankers, compressors, generators, and evaluators into replaceable modules. The cost is higher engineering complexity: graph construction, incremental updates, module interfaces, and evaluation all become more difficult. In practice, Advanced RAG improves recall, Graph RAG improves relational understanding, and Agentic RAG decides when to call each path.",
    },
    {
        "no": "21",
        "title": "Agentic RAG / Agentic RAG",
        "cn": "第五部分进入 Agentic RAG。普通 RAG 通常是固定的检索生成链路，而 Agentic RAG 把它升级为可规划、可观察、可重试、可兜底的动态流程。这里的关键词是“决策”。系统不再每次机械地检索同一个知识库，而是先判断任务类型，再决定是否检索、查哪个知识源、是否调用工具、是否需要多轮检索、证据是否足够，以及什么时候应该拒答或转人工。它适合复杂业务流程，但也会带来成本、延迟和可控性挑战。",
        "en": "The fifth section introduces Agentic RAG. Standard RAG usually follows a fixed retrieval-generation chain, while Agentic RAG upgrades it into a dynamic flow that can plan, observe, retry, and fall back. The keyword is decision-making. The system no longer mechanically retrieves from the same knowledge base every time. It first judges the task type, then decides whether to retrieve, which source to query, whether to call tools, whether multiple rounds are needed, whether the evidence is sufficient, and when to refuse or hand off to a human. It is suitable for complex business workflows, but it also introduces cost, latency, and control challenges.",
    },
    {
        "no": "22",
        "title": "Classic vs Agentic / 普通 RAG 与 Agentic RAG",
        "cn": "Agentic RAG 的差异不在于多一个 Agent，而在于检索链路由谁决策。普通 RAG 是固定线性流程，通常每次都检索，多数只查一次知识库，错误恢复能力较弱，适合单跳知识问答。Agentic RAG 是动态决策流程，Agent 可以判断是否检索、检索几次、是否跨多个来源、是否调用数据库、API、搜索或计算工具，也可以重写查询、重试、反思和兜底。它适合复杂任务、多系统协作和业务流程，但成本和延迟更高，所以必须设置更严格的边界。",
        "en": "The difference in Agentic RAG is not that we add one more Agent; it is who decides the retrieval path. Classic RAG follows a fixed linear flow. It usually retrieves every time, often queries the knowledge base once, has weak error recovery, and works best for single-hop knowledge QA. Agentic RAG uses a dynamic decision flow. The Agent can decide whether to retrieve, how many times to retrieve, whether to use multiple sources, whether to call databases, APIs, search, or calculation tools, and whether to rewrite, retry, reflect, or fall back. It is better for complex tasks, multi-system collaboration, and business workflows, but it has higher cost and latency, so stronger boundaries are required.",
    },
    {
        "no": "23",
        "title": "Agent Loop / Agent 决策循环",
        "cn": "这一页展示 Agentic RAG 的循环。核心不是盲目多检索，而是在观察证据后再决策。系统先理解问题，规划任务，再选择工具或知识源。拿到结果后，它要观察证据是否相关、是否充分、是否冲突。如果证据足够，就生成答案并自检；如果不足，就改写查询、换工具、继续检索，或者触发兜底。这里必须设置最大重试次数、权限校验、参数验证和人工兜底，否则动态流程可能陷入循环、越权调用工具，或者产生难以审计的行为。",
        "en": "This slide shows the Agentic RAG loop. The core is not retrieving more blindly, but deciding after observing evidence. The system first understands the question, plans the task, and chooses a tool or knowledge source. After receiving results, it evaluates whether the evidence is relevant, sufficient, or conflicting. If the evidence is enough, it generates and checks the answer. If not, it may rewrite the query, switch tools, retrieve again, or trigger fallback. This flow must include retry limits, permission checks, parameter validation, and human handoff. Otherwise, the dynamic process may loop indefinitely, call tools without authorization, or produce behavior that is hard to audit.",
    },
    {
        "no": "24",
        "title": "Agent Roles / Agent 角色分工",
        "cn": "Agentic RAG 要灵活又可控，就必须明确角色边界。路由器判断问题属于知识库问答、工具查询、闲聊、拒答还是人工处理。规划器把复杂问题拆成子问题。检索控制器决定查哪个知识库、用什么查询词、查几次。工具调用者通过 Function Calling 查询订单、数据库、搜索引擎或计算器。证据评估器判断材料是否相关、充分或冲突。生成控制器组织证据并约束模型回答。反思器检查答案是否忠实完整。兜底处理器在证据不足或风险较高时澄清、拒答或转人工。动态越强，控制点越要前置。",
        "en": "For Agentic RAG to be flexible and controllable, role boundaries must be explicit. The router classifies the question as knowledge-base QA, tool query, casual chat, refusal, or human handling. The planner breaks complex questions into subquestions. The retrieval controller decides which knowledge base to query, what query terms to use, and how many times to search. The tool caller uses function calling to query orders, databases, search engines, or calculators. The evidence evaluator checks whether evidence is relevant, sufficient, or conflicting. The generation controller organizes evidence and constrains the model answer. The reflector checks whether the answer is faithful and complete. The fallback handler clarifies, refuses, or escalates when evidence is insufficient or risky. The more dynamic the system is, the earlier control points must be designed.",
    },
    {
        "no": "25",
        "title": "Complex Case / 复杂客服案例",
        "cn": "这个客服案例说明多源、多步 Agentic RAG 的价值。用户问：订单 A1001 还在运输中，但耳机坏了，现在能退款吗？系统不能只查 FAQ，也不能直接承诺退款。第一步识别订单号，并调用订单工具确认状态。第二步识别售后意图，检索故障与退款政策。第三步组合证据，把订单状态、签收时间和条款放在一起判断。第四步发现缺口，如果缺少签收时间或故障证明，就请求补充。第五步生成答案，标明订单工具和政策依据。最后还要反思，如果证据不可靠，就不能承诺，应该转人工或说明限制。",
        "en": "This customer-service case shows the value of multi-source, multi-step Agentic RAG. The user asks: Order A1001 is still in transit, but the headphones are broken. Can I get a refund now? The system should not only search FAQs, and it should not promise a refund directly. First, it identifies the order ID and calls the order lookup tool to confirm status. Second, it identifies the after-sales intent and retrieves fault and refund policies. Third, it combines evidence: order status, delivery time, and policy terms. Fourth, it detects gaps. If delivery time or fault proof is missing, it asks for clarification. Fifth, it generates an answer with references to the order tool and policy. Finally, it reflects: if evidence is unreliable, it should not make a promise, and should escalate or explain the limitation.",
    },
    {
        "no": "26",
        "title": "Benefits and Risks / 优势与风险",
        "cn": "Agentic RAG 的优势是能处理复杂、多步骤、多来源任务；能根据问题类型选择知识库、工具或人工兜底；能在检索失败后主动改写查询和重试；也能对检索结果和最终答案做自检，所以更适合企业业务系统集成。但风险同样明显：多次调用 LLM 和检索器会增加成本和延迟；Agent 可能做错决策，比如不该调用工具时调用工具；多轮流程可能陷入循环；Tool Use 会带来权限和参数安全问题；反思机制本身也不绝对可靠。所以上线评估不能只看回答质量，还要看工具安全、循环控制和恢复行为。",
        "en": "Agentic RAG is useful because it can handle complex, multi-step, multi-source tasks. It can choose a knowledge base, tool, or human fallback based on question type. It can rewrite queries and retry after retrieval failure. It can also self-check retrieval results and final answers, making it better suited for enterprise business-system integration. But the risks are equally clear. Repeated LLM and retriever calls increase cost and latency. The Agent may make wrong decisions, such as calling tools unnecessarily. Multi-round flows may loop. Tool use introduces permission and parameter-safety risks. Reflection is not fully reliable. Therefore, production evaluation should not only measure answer quality; it must also evaluate tool safety, loop control, and recovery behavior.",
    },
    {
        "no": "27",
        "title": "Applications, Challenges, and Evaluation / 应用、挑战与评估",
        "cn": "第六部分把讨论落到应用、挑战和评估。RAG 的工程价值最终不取决于架构图是否漂亮，而取决于它能不能在真实场景里稳定、准确、可控地回答问题。典型场景包括企业知识库、智能客服、法律医疗、金融研究、代码助手和多模态检索。与此同时，我们也要面对检索不到、上下文噪声、切分不合理、数据安全、模型仍会幻觉等问题。最后，评估体系必须能区分检索质量和生成质量，否则很难知道问题应该在哪一层修。",
        "en": "The sixth section brings the discussion to applications, challenges, and evaluation. The engineering value of RAG does not depend on whether the architecture diagram looks elegant. It depends on whether the system can answer real questions stably, accurately, and controllably. Typical scenarios include enterprise knowledge bases, customer service, law and medicine, financial research, code assistants, and multimodal retrieval. At the same time, we must handle missed retrieval, noisy context, poor chunking, data security, and continued hallucination. Finally, the evaluation system must separate retrieval quality from generation quality; otherwise, we cannot know which layer needs to be fixed.",
    },
    {
        "no": "28",
        "title": "Use Cases / 典型场景",
        "cn": "RAG 最适合知识私有、变化快、且回答必须有依据的场景。企业知识库里，制度、流程和产品文档经常更新，重点是权限、版本和引用准确性。智能客服可以结合 FAQ、售后政策和订单工具，但不能编造承诺，需要人工兜底。法律和医疗必须保留引用并由专业人员复核。金融研报关注时效性、合规性和来源可信度。代码助手需要理解项目结构、版本和依赖。多模态检索则要处理图纸、截图、表格和扫描件。无论场景如何，共同底线都是权限控制、版本管理、引用准确和人工复核。",
        "en": "RAG is strongest when knowledge is private, changes quickly, and answers require evidence. In enterprise knowledge bases, policies, processes, and product documents update frequently, so permissions, versions, and citation accuracy matter. Customer service can combine FAQs, after-sales policies, and order tools, but it must not fabricate commitments and needs human fallback. Law and medicine require citations and professional review. Financial research depends on timeliness, compliance, and source reliability. Code assistants need to understand project structure, versions, and dependencies. Multimodal retrieval must handle drawings, screenshots, tables, and scanned files. Across all scenarios, the shared baseline is permission control, version management, accurate citation, and human review.",
    },
    {
        "no": "29",
        "title": "Advantages / RAG 的优势",
        "cn": "RAG 的优势来自证据外置，而不是简单让模型变大。第一，它可以降低幻觉，因为模型基于外部证据回答，而不是纯靠参数记忆。第二，它支持知识更新，更新文档和索引通常比重新训练模型更快。第三，它能接入私有数据，把企业内部文档作为外部知识库。第四，它提高可追溯性，可以返回来源、段落、链接和时间戳。第五，它降低微调成本，知识变化不必每次训练模型。第六，它便于权限控制，在检索阶段按用户身份过滤数据。也就是说，RAG 给企业带来的是一组可操作的工程抓手。",
        "en": "RAG’s advantages come from externalizing evidence, not simply making the model larger. First, it reduces hallucination because the model answers from external evidence rather than only from parameter memory. Second, it supports knowledge updates: updating documents and indexes is usually faster than retraining a model. Third, it connects private data by turning internal enterprise documents into external knowledge bases. Fourth, it improves traceability by returning source, section, link, and timestamp. Fifth, it lowers fine-tuning cost because knowledge changes do not require retraining every time. Sixth, it supports permission control by filtering data during retrieval based on user identity. In short, RAG gives enterprises practical engineering handles.",
    },
    {
        "no": "30",
        "title": "Challenges / RAG 的挑战",
        "cn": "RAG 的挑战必须按链路定位。检索不到正确文档时，要优化 chunk、Embedding、混合检索和查询改写。噪声太多时，要用 Rerank、相似度阈值和上下文压缩。Chunk 不合理时，要按结构切分、使用重叠窗口或语义切分。上下文过长时，要调 Top-K、做摘要压缩或分层检索。多跳问题困难时，需要问题分解、迭代检索或 Agentic RAG。数据安全问题要靠权限标签、租户隔离和审计日志。即使有证据，模型仍可能幻觉，所以还需要严格 Prompt、引用校验、反思和拒答机制。不同症状对应不同缓解方式。",
        "en": "RAG challenges must be located by pipeline stage. If the correct documents are not retrieved, we should improve chunks, embeddings, hybrid search, and query rewriting. If there is too much noise, we need reranking, similarity thresholds, and context compression. If chunks are poor, we need structure-aware splitting, overlap windows, or semantic splitting. If context is too long, we should tune Top-K, summarize, or use hierarchical retrieval. If multi-hop questions are difficult, we need question decomposition, iterative retrieval, or Agentic RAG. Data security requires permission tags, tenant isolation, and audit logs. Even with evidence, the model may still hallucinate, so strict prompts, citation checks, reflection, and refusal mechanisms are needed. Different symptoms require different mitigations.",
    },
    {
        "no": "31",
        "title": "Evaluation / 评估体系",
        "cn": "RAG 评估必须拆开检索质量和生成质量。检索侧可以看 Hit Rate@K，也就是 Top-K 里是否包含正确文档；看 Recall@K 或 Context Recall，判断必要上下文召回了多少；看 Precision@K 或 Context Precision，判断检索内容有多少真正相关；也可以看 MRR 或 nDCG，衡量正确文档排序位置和相关性等级。生成侧要看 Faithfulness，也就是答案是否忠于上下文；看 Answer Relevance 或 Correctness，判断是否回答问题、事实是否正确；还要看 Citation Accuracy，引用是否真的支持答案。线上则结合自动评估、人工抽检、用户反馈、成本和延迟监控。",
        "en": "RAG evaluation must separate retrieval quality from generation quality. On the retrieval side, we can measure Hit Rate@K, meaning whether the correct document appears in Top-K. We can measure Recall@K or Context Recall, meaning how much required context was retrieved. We can measure Precision@K or Context Precision, meaning how much retrieved content is truly relevant. We can also use MRR or nDCG to measure the ranking of correct documents and graded relevance. On the generation side, we evaluate Faithfulness, or whether the answer is grounded in the context; Answer Relevance or Correctness, or whether the answer addresses the question and is factually correct; and Citation Accuracy, or whether cited sources truly support the answer. Online monitoring combines automated evaluation, human sampling, user feedback, cost, and latency.",
    },
    {
        "no": "32",
        "title": "Relationship with Other Techniques / 与其他技术的关系",
        "cn": "RAG 不是要替代所有技术，而是和微调、长上下文、Prompt、知识图谱和 Agent 互补。RAG 提供外部事实依据。Fine-tuning 更适合改变模型行为、输出格式和领域能力，但不适合频繁更新事实知识。长上下文能容纳更多材料，但仍然需要选择和组织，否则会成本高、噪声多。Prompt Engineering 决定模型如何使用资料。Knowledge Graph 提供结构化关系和全局分析。Agent 则负责规划和工具调用，可以把 RAG 当作一个工具来使用。清晰分工能避免把所有问题都强行塞给一种技术路线。",
        "en": "RAG is not meant to replace all other techniques. It complements fine-tuning, long context, prompt engineering, knowledge graphs, and Agents. RAG provides external factual evidence. Fine-tuning is better for changing model behavior, output format, and domain capability, but it is not ideal for frequently updated facts. Long context can hold more material, but it still needs selection and organization; otherwise, cost and noise increase. Prompt engineering defines how the model should use evidence. Knowledge graphs provide structured relationships and global analysis. Agents handle planning and tool use, and they can call RAG as one tool. Clear division of labor prevents every problem from being forced into one technical route.",
    },
    {
        "no": "33",
        "title": "Future / 未来趋势",
        "cn": "最后一页总结未来趋势。RAG 会从固定问答增强，走向动态、结构化、多模态和自优化。Agentic RAG 会根据任务复杂度选择检索、工具和验证路径。Graph RAG 会用实体和关系图增强关系推理。Multimodal RAG 会跨文本、图像、音频和视频检索生成。Real-time RAG 会连接实时数据库、搜索、业务 API 和消息流。Personalized RAG 会结合身份、历史、角色和业务上下文。Permission-aware RAG 会把权限控制贯穿索引、检索、生成和引用。On-device RAG 会支持本地笔记和文件助手。自动化评估会发现低召回、调参、生成测试集并做回归比较。最终目标不是跑通 Demo，而是稳定、准确、可控地回答真实问题。",
        "en": "The final slide summarizes future trends. RAG will move from fixed question answering toward dynamic, structured, multimodal, and self-optimizing systems. Agentic RAG will choose retrieval, tools, and verification paths based on task complexity. Graph RAG will use entity and relation graphs to improve relationship reasoning. Multimodal RAG will retrieve and generate across text, images, audio, and video. Real-time RAG will connect live databases, search engines, business APIs, and streams. Personalized RAG will use identity, history, role, and business context. Permission-aware RAG will apply access control across indexing, retrieval, generation, and citation. On-device RAG will support local notes and file assistants. Automated evaluation will find low recall, tune parameters, create test sets, and run regression comparisons. The ultimate goal is not just to run a demo, but to answer real questions stably, accurately, and controllably.",
    },
]


def set_font(run, size=None, bold=False, color=None, east_asia="STHeiti"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_paragraph(doc, text, style=None, size=10.5, bold=False, color=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def build_markdown():
    lines = [
        "# RAG 中英双语展示讲稿",
        "",
        "说明：本稿按英文版 PPT 的 33 页逐页对应，每页提供中文讲稿与英文讲稿。",
        "",
    ]
    for s in slides:
        lines.extend([
            f"## Slide {s['no']} - {s['title']}",
            "",
            "### 中文讲稿",
            s["cn"],
            "",
            "### English Script",
            s["en"],
            "",
        ])
    return "\n".join(lines)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "RAG bilingual speaker script"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=9, color="666666")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("RAG 中英双语展示讲稿")
    set_font(title_run, size=24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("逐页对应英文版 PPT，共 33 页。每页包含中文讲稿与英文讲稿，可直接用于排练和展示。")
    set_font(subtitle_run, size=11, color="666666")

    meta = add_paragraph(doc, "对应文件：RAG_Technical_Review_EN.pptx；内容原则：不新增主题，不改变页序，按页展开讲解。", size=10, color="666666", after=12)
    shade_paragraph(meta, "F4F6F9")

    for s in slides:
        h = doc.add_heading(f"Slide {s['no']}  {s['title']}", level=1)
        for run in h.runs:
            set_font(run, size=16, bold=True, color="2E74B5")

        cn_h = doc.add_heading("中文讲稿", level=2)
        for run in cn_h.runs:
            set_font(run, size=13, bold=True, color="1F4D78")
        add_paragraph(doc, s["cn"], size=10.5, after=8)

        en_h = doc.add_heading("English Script", level=2)
        for run in en_h.runs:
            set_font(run, size=13, bold=True, color="1F4D78")
        add_paragraph(doc, s["en"], size=10.5, after=10)

    doc.save(DOCX)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    MD.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(DOCX)
    print(MD)


if __name__ == "__main__":
    main()
