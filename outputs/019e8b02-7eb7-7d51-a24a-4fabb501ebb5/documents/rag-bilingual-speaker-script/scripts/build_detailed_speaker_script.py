from pathlib import Path
import importlib.util

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_SCRIPT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/scripts/build_speaker_script.py")
OUT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/output")
DOCX = OUT / "RAG_中英双语展示讲稿_详细版.docx"
MD = OUT / "RAG_中英双语展示讲稿_详细版.md"


spec = importlib.util.spec_from_file_location("base_speaker_script", BASE_SCRIPT)
base = importlib.util.module_from_spec(spec)
spec.loader.exec_module(base)
slides = base.slides


details = {
    "01": {
        "cn": [
            "这一页还可以先帮听众建立一个直觉：大模型本身像一个很强的语言和推理引擎，但它不是实时数据库，也不是企业内部知识库。RAG 做的事情，是在模型回答前增加一个“查证层”，让模型先看到和问题相关的证据，再组织答案。",
            "可以举一个简单例子：如果用户问“最新报销政策里差旅住宿标准是多少”，单靠模型参数可能会答错，因为政策可能刚更新；RAG 则先检索最新制度文件，再让模型基于那一段制度回答。这样答案不仅更准确，还能说明依据来自哪份文件、哪一节、哪个时间版本。",
            "开场时建议强调本次汇报的主线：先讲 RAG 为什么出现，再讲它怎么工作，再讲工程系统如何落地，最后讨论 Agentic RAG 和评估。这样听众会知道后面的技术点不是散的，而是围绕“如何让模型可靠使用外部证据”展开。",
        ],
        "en": [
            "At this point, it is useful to give the audience an intuition: a large language model is a powerful language and reasoning engine, but it is not a real-time database and it is not automatically connected to an enterprise knowledge base. RAG adds an evidence-checking layer before generation, so the model can see relevant evidence first and then formulate the answer.",
            "A simple example is an employee asking, “What is the latest hotel allowance in the travel reimbursement policy?” If the policy was updated recently, relying only on model parameters may produce a wrong answer. With RAG, the system first retrieves the latest policy document, then asks the model to answer based on the relevant section. The answer becomes more accurate and can point to the document, section, and version.",
            "For the opening, emphasize the storyline of the talk: why RAG exists, how it works, how it is engineered in production, and why Agentic RAG and evaluation matter. This helps the audience see that all technical details serve one theme: making the model use external evidence reliably.",
        ],
    },
    "02": {
        "cn": [
            "这一页可以把“动机”讲得更业务化一些。企业不是单纯追求模型能不能聊天，而是关心答案能否被信任、能否被复核、能否适应知识变化。RAG 正是围绕这些可靠性需求出现的。",
            "只靠参数记忆有一个天然矛盾：模型越大，记住的通用知识越多，但企业越关心的往往越是新知识、私有知识和受权限控制的知识。这些知识不适合直接放进训练数据，也不一定允许被所有用户看到。",
            "因此，本部分的关键结论是：RAG 不是为了替代大模型，而是为了补足大模型在事实更新、私有知识接入和可追溯性方面的短板。后面所有链路设计，都是围绕这三个短板展开。",
        ],
        "en": [
            "This slide can be framed in business terms. Enterprises do not only care whether a model can chat fluently. They care whether the answer can be trusted, reviewed, and adapted when knowledge changes. RAG emerged because of these reliability requirements.",
            "Parameter memory has a natural limitation: the larger the model, the more general knowledge it may contain, but enterprises often care most about new knowledge, private knowledge, and permission-controlled knowledge. Such knowledge is not always suitable for training data, and it should not be visible to every user.",
            "The key conclusion of this section is that RAG does not replace the language model. It complements the model by solving problems around factual updates, private knowledge access, and traceability. The pipeline designs in later slides all serve these goals.",
        ],
    },
    "03": {
        "cn": [
            "定义这一页建议按流程逐格讲。用户问题进入系统后，第一步不是直接让模型回答，而是先把问题转化为检索请求。检索对象可以是文档库，也可以是数据库、搜索引擎、业务 API 或知识图谱。",
            "检索到的内容不会原样全部交给模型，而是要经过相关性筛选、Top-K 选择、过滤和必要的压缩。随后系统把证据片段、用户问题和回答约束一起构造成 Prompt。模型看到的是一个被组织过的上下文，而不是杂乱的全文。",
            "最后的引用返回非常重要。没有 citation 的 RAG 只能算“看过资料后回答”，但有了 source、chunk、link 或时间戳，答案才具备审计能力。尤其在企业场景里，可追溯性往往和准确性同样重要。",
        ],
        "en": [
            "For this definition slide, walk through the flow step by step. When a user question enters the system, the first step is not to let the model answer directly. The system first converts the question into a retrieval request. The retrieval source may be a document repository, database, search engine, business API, or knowledge graph.",
            "The retrieved content is not sent to the model in raw form. It goes through relevance selection, Top-K filtering, permission checks, and sometimes compression. Then the evidence chunks, user question, and answer constraints are assembled into a prompt. The model receives organized context rather than an unfiltered document dump.",
            "Citation return is critical. Without citations, RAG only means that the model answered after seeing some material. With source IDs, chunk IDs, links, or timestamps, the answer becomes auditable. In enterprise settings, traceability is often just as important as accuracy.",
        ],
    },
    "04": {
        "cn": [
            "讲五类问题时，可以把它们和真实风险对应起来。知识滞后会导致回答旧政策；幻觉会导致模型编造不存在的条款；私有知识缺失会导致模型不了解企业内部流程；长上下文成本会让系统又慢又贵；依据不透明会让业务团队无法复核。",
            "这些问题在开放聊天场景里可能只是体验问题，但在企业场景里会变成责任问题。比如客服错误承诺退款、法律助手引用错误条款、医疗问答给出无依据建议，都会带来业务和合规风险。",
            "所以 RAG 的价值不是简单提高回答“看起来专业”的程度，而是降低业务风险。它把答案绑定到可检查的证据上，让知识更新、权限控制和结果审计成为系统能力。",
        ],
        "en": [
            "When explaining the five problems, connect each one to a real risk. Stale knowledge may lead to outdated policy answers. Hallucination may cause the model to invent nonexistent clauses. Missing private knowledge means the model does not understand internal processes. Long context makes the system slow and expensive. Opaque evidence makes review difficult.",
            "In open-ended chat, these may be user-experience issues. In enterprise settings, they become accountability issues. A customer-service assistant may promise a refund incorrectly, a legal assistant may cite the wrong clause, or a medical assistant may provide advice without evidence. These failures carry business and compliance risks.",
            "Therefore, RAG is not just about making answers sound more professional. It is about reducing operational risk by grounding answers in checkable evidence. It turns knowledge updates, permission control, and answer auditability into system capabilities.",
        ],
    },
    "05": {
        "cn": [
            "这里可以用“脑内知识”和“外部资料库”做类比。参数知识像模型已经学会的语言能力、常识和推理模式；非参数知识像一本随时可以更新、可以查目录、可以标注来源的资料库。",
            "参数知识的优势是速度快、泛化强，不需要每次都查资料；但它的问题是不可控、不可直接更新、也不知道具体来源。非参数知识正好相反，它可以被替换、删除、加权限、加版本，也可以记录每次命中的证据。",
            "RAG 的设计重点就是让两者各做擅长的事。不要要求参数记忆承担所有事实知识，也不要指望检索系统自己会推理和表达。好的 RAG 系统，是让检索负责事实，模型负责理解、归纳和生成。",
        ],
        "en": [
            "A useful analogy is “knowledge inside the brain” versus an “external reference library.” Parametric knowledge is the model’s learned language ability, common sense, and reasoning patterns. Non-parametric knowledge is like an external library that can be updated, searched, versioned, and cited.",
            "Parametric knowledge is fast and general. It does not require retrieval every time. But it is hard to control, hard to update directly, and not tied to a specific source. Non-parametric knowledge is the opposite: it can be replaced, deleted, permissioned, versioned, and logged as evidence.",
            "The design principle of RAG is to let each part do what it is good at. We should not force parameter memory to carry all factual knowledge, and we should not expect the retrieval system to reason and communicate by itself. A good RAG system lets retrieval handle facts and lets the model handle understanding, synthesis, and generation.",
        ],
    },
    "06": {
        "cn": [
            "核心链路这一页可以先讲一个重要判断：RAG 的失败通常是链路型失败，而不是单点失败。一个答案不可靠，可能不是因为模型差，而是因为文档没解析出来、chunk 切错、检索没召回、排序不合理或 Prompt 没约束。",
            "离线索引像是在建图书馆的目录系统。目录建得不好，再聪明的读者也找不到书。在线问答像是在规定时间内帮用户找到相关章节并写出回答，要求速度、准确性和可引用性同时满足。",
            "因此，后面讲每个环节时要保持一个判断标准：这个环节如何影响最终答案？它影响的是召回率、上下文质量、生成忠实性，还是引用可追溯性？",
        ],
        "en": [
            "For the core pipeline slide, start with an important observation: RAG failures are usually pipeline failures, not single-point failures. An unreliable answer may not mean the model is weak. The cause may be failed parsing, poor chunking, missed recall, bad ranking, or insufficient prompt constraints.",
            "Offline indexing is like building the catalog system of a library. If the catalog is poor, even a very smart reader cannot find the right book. Online QA is like helping the user find the relevant sections and write an answer under time pressure, while still maintaining accuracy and citations.",
            "When discussing each later component, keep asking: how does this stage affect the final answer? Does it affect recall, context quality, generation faithfulness, or citation traceability?",
        ],
    },
    "07": {
        "cn": [
            "端到端链路可以按“数据变成证据，证据变成答案”来讲。采集和清洗解决的是资料能不能被机器正确读取；切分和向量化解决的是资料能不能被检索；索引和检索解决的是用户提问时能不能找到相关内容；生成和引用解决的是能不能把证据转成可信答案。",
            "一个常见误区是只盯着最后的 LLM。比如答案错了就换更大的模型，但如果错误来自 PDF 表格解析失败，换模型并不能根治。如果问题是 chunk 把定义和例外条款切开了，模型也只能看到不完整证据。",
            "所以生产级 RAG 的优化要做链路诊断：先看正确文档是否入库，再看是否被召回，再看是否排在前面，再看模型是否正确使用证据。这样的排查顺序比盲目调 Prompt 更有效。",
        ],
        "en": [
            "The end-to-end pipeline can be explained as “data becomes evidence, and evidence becomes an answer.” Ingestion and cleaning determine whether materials are machine-readable. Chunking and embedding determine whether materials can be retrieved. Indexing and retrieval determine whether relevant content can be found when a question arrives. Generation and citation determine whether evidence becomes a trustworthy answer.",
            "A common mistake is focusing only on the final LLM. If an answer is wrong, teams may immediately try a larger model. But if the root cause is failed PDF table parsing, changing the model will not solve the problem. If chunking separates a rule from its exception, the model only sees incomplete evidence.",
            "Production RAG optimization should therefore follow pipeline diagnosis: first check whether the correct document is indexed, then whether it is retrieved, then whether it is ranked high enough, and finally whether the model uses the evidence correctly. This is more effective than blindly tuning prompts.",
        ],
    },
    "08": {
        "cn": [
            "离线索引的难点在于原始知识通常很脏。PDF 可能有页眉页脚、分栏、表格和脚注；网页可能有导航栏和广告；数据库字段可能缺少说明；代码仓库还涉及依赖、函数调用和版本。",
            "清洗阶段不是简单删噪声，而是要保留对回答有用的结构。例如标题层级、表格行列、章节编号、文档版本、更新时间、权限标签，都可能在后续检索、过滤和引用中发挥作用。",
            "向量存储时，metadata 和向量本身同样重要。没有 metadata，就很难做权限过滤、版本过滤、来源展示和审计追踪。可以强调：RAG 的证据库不是一堆文本向量，而是“文本片段 + 结构信息 + 权限信息 + 来源信息”的组合。",
        ],
        "en": [
            "The difficulty of offline indexing is that raw knowledge is usually messy. PDFs may contain headers, footers, columns, tables, and footnotes. Web pages may include navigation bars and advertisements. Database fields may lack descriptions. Code repositories involve dependencies, function calls, and versions.",
            "Cleaning is not simply deleting noise. It is also about preserving useful structure. Headings, table rows and columns, section numbers, document versions, update timestamps, and permission labels may all matter later for retrieval, filtering, and citation.",
            "When storing vectors, metadata is as important as the vectors themselves. Without metadata, it becomes difficult to filter by permission, select versions, show sources, or audit answers. A RAG evidence base is not just a pile of text vectors. It is a combination of text chunks, structural information, permission information, and source information.",
        ],
    },
    "09": {
        "cn": [
            "Chunking 可以结合一个例子讲：如果公司政策里前半段写“员工可以报销”，后半段写“但高铁商务座除外”，如果切分把两句话分到不同 chunk，模型只看到前半段就可能给出错误承诺。",
            "切得太大也有问题。一个 chunk 里包含多个无关主题时，检索虽然命中了，但模型会被噪声干扰，甚至引用到不相关段落。切得太小则缺少上下文，尤其法律、财务、技术文档中很多条款依赖上下文。",
            "实践中常见策略是混合使用：先按章节结构切，再对长段落递归切分；对表格、代码、FAQ 做专门处理；对关键文档保留标题路径和前后窗口。好的 chunk 应该让检索容易命中，也让模型容易理解。",
        ],
        "en": [
            "Chunking can be explained with a policy example. Suppose the first sentence says “employees may claim reimbursement,” while a later sentence says “except for business-class high-speed rail.” If chunking separates the rule and the exception, the model may only see the first part and make a wrong promise.",
            "Chunks that are too large also create problems. If one chunk contains many unrelated topics, retrieval may technically hit the document, but the model may be distracted by noise or cite irrelevant paragraphs. Chunks that are too small lack context, which is especially dangerous in legal, financial, and technical documents where clauses depend on surrounding text.",
            "In practice, teams often combine strategies: split by section structure first, recursively split long passages, treat tables, code, and FAQs specially, and preserve heading paths and neighboring windows for important documents. A good chunk should be easy to retrieve and easy for the model to understand.",
        ],
    },
    "10": {
        "cn": [
            "Embedding 模型的作用是把文本变成可比较的语义坐标。用户问“退款”，文档里写“售后退货”；用户问“报销”，文档里写“费用申请”，语义模型应该能把这些表达拉近。",
            "但 Embedding 不是万能的。它可能对产品编号、合同编号、金额、日期、代码符号不敏感，也可能对行业缩写理解不足。因此在生产中，常常需要把 Dense Retrieval 和 BM25 结合，避免语义模型漏掉硬关键词。",
            "选择 Embedding 时要做领域评测，而不是只看公开榜单。可以准备一组企业真实问题，检查正确文档能否进入 Top-K；同时比较中英文、缩写、口语化问题、长问题和多跳问题的表现。",
        ],
        "en": [
            "An embedding model converts text into semantic coordinates that can be compared. If the user asks about “refund,” while the document says “after-sales return,” or the user asks about “reimbursement,” while the policy says “expense claim,” the embedding model should map these expressions close to each other.",
            "But embeddings are not magic. They may be weak on product IDs, contract numbers, amounts, dates, and code symbols. They may also misunderstand domain abbreviations. This is why production systems often combine dense retrieval with BM25, so semantic retrieval does not miss hard lexical constraints.",
            "Embedding selection should be based on domain evaluation, not only public leaderboards. Prepare real enterprise questions and check whether the correct document enters Top-K. Also compare performance on Chinese and English, abbreviations, conversational queries, long questions, and multi-hop questions.",
        ],
    },
    "11": {
        "cn": [
            "向量数据库的核心问题是规模化检索。小数据量时可以逐个算距离，但百万、千万甚至更大规模时，精确搜索会变得太慢，所以需要 ANN 近似最近邻索引。",
            "不同索引结构适合不同场景。HNSW 适合高召回低延迟，但内存成本高；IVF 更适合可调参数的大规模场景；PQ 牺牲一点精度换存储成本；DiskANN 则把部分能力转向磁盘，适合更大规模。",
            "除了索引算法，还要看工程能力：是否支持 metadata filter，是否支持增量更新，是否支持多租户隔离，是否有备份恢复，是否方便和权限系统打通。向量库选型不是算法题，而是检索质量和生产运维的平衡题。",
        ],
        "en": [
            "The core issue for vector databases is scalable retrieval. With small datasets, exact distance calculation is possible. But at millions or tens of millions of vectors, exact search becomes too slow, so approximate nearest neighbor, or ANN, indexes are needed.",
            "Different index structures fit different scenarios. HNSW is good for high recall and low latency, but it requires more memory. IVF is suitable for large-scale settings with tunable parameters. PQ trades some accuracy for lower storage cost. DiskANN shifts part of the workload to disk and is useful at very large scale.",
            "Beyond index algorithms, engineering capabilities matter: metadata filtering, incremental updates, multi-tenant isolation, backup and recovery, and integration with permission systems. Choosing a vector database is not only an algorithm decision; it is a balance between retrieval quality and production operations.",
        ],
    },
    "12": {
        "cn": [
            "BM25 和 Dense Retrieval 的差异可以用“精确词”和“语义意图”解释。BM25 看到的是词面重合，所以对 SKU、政策编号、函数名、错误码很有效。Dense Retrieval 看到的是语义接近，所以对同义表达、自然语言描述和跨语言问题更有效。",
            "混合检索的价值在于互补。比如用户问“iPhone 15 Pro 退货时间”，BM25 能抓住产品型号和“退货”，Dense 能理解“退货时间”和“售后政策”的语义关系。两者合并后，再通过 rerank 排序，通常比单一方法更稳。",
            "这里也可以提醒听众：检索不是一个模型越新越好的单变量问题，而是要根据数据类型设计。代码、合同、订单、FAQ、知识库文章，对检索方式的要求都不同。",
        ],
        "en": [
            "The difference between BM25 and dense retrieval can be explained as exact wording versus semantic intent. BM25 relies on lexical overlap, so it is strong for SKUs, policy IDs, function names, and error codes. Dense retrieval relies on semantic similarity, so it is strong for paraphrases, natural-language descriptions, and cross-lingual questions.",
            "Hybrid search is valuable because the two methods complement each other. If the user asks about “iPhone 15 Pro return period,” BM25 captures the product model and the word “return,” while dense retrieval connects “return period” with “after-sales policy.” Combined retrieval, followed by reranking, is usually more robust than either method alone.",
            "This is also a good place to remind the audience that retrieval is not a single-variable problem where the newest model always wins. Retrieval design depends on data type. Code, contracts, orders, FAQs, and knowledge-base articles each require different retrieval strategies.",
        ],
    },
    "13": {
        "cn": [
            "重排序可以理解为第二轮更精细的判断。第一轮检索为了不漏，可能召回很多候选；Reranker 再把用户问题和候选片段一起看，判断哪几个真正适合放进上下文。",
            "上下文组织也很关键。不是把 Top-K 按分数堆进去就行，而是要去重复、保持逻辑顺序、保留标题路径、必要时压缩长片段，并把来源信息放在模型可见的位置。否则模型可能看到证据，却不知道证据来自哪里、适用范围是什么。",
            "Prompt 约束要写清楚边界：只能根据资料回答；证据不足时说明不确定；引用要对应到具体来源；如果资料冲突，要指出冲突而不是强行合并。RAG 的 Prompt 不是让模型更会发挥，而是让模型更受证据约束。",
        ],
        "en": [
            "Reranking can be understood as a more precise second-stage judgment. The first retrieval stage tries not to miss anything, so it may return many candidates. The reranker then reads the user question and each candidate chunk together to decide which ones truly deserve to enter the context.",
            "Context assembly is equally important. We should not simply dump Top-K chunks by score. The system should deduplicate, preserve logical order, keep heading paths, compress long chunks when needed, and place source information where the model can see it. Otherwise, the model may see evidence without understanding where it comes from or when it applies.",
            "Prompt constraints must define boundaries clearly: answer only from supplied evidence, say when evidence is insufficient, cite specific sources, and highlight conflicts instead of forcing a single answer. In RAG, the prompt is not meant to make the model more imaginative; it is meant to make the model more evidence-constrained.",
        ],
    },
    "14": {
        "cn": [
            "工程实现部分可以先区分 Demo 和生产。Demo 通常只关心能不能跑通一条链路；生产系统要考虑数据同步、权限、审计、故障恢复、成本、延迟和持续评估。",
            "真实系统中还会有很多边界问题：文档更新后索引是否同步？员工离职后权限是否立即生效？用户问的问题是否涉及敏感信息？答案引用的版本是否过期？这些都不属于单纯模型能力，而属于系统工程。",
            "因此，工程实现的核心不是堆工具，而是把 RAG 拆成可观测、可替换、可评估的层。只有这样，系统出问题时才能定位到具体层，而不是在模型、检索、数据之间来回猜。",
        ],
        "en": [
            "In the engineering section, first distinguish a demo from production. A demo usually only proves that one flow can run. A production system must handle data synchronization, permissions, audit logs, recovery, cost, latency, and continuous evaluation.",
            "Real systems also face many boundary issues. Is the index updated after a document changes? Does permission change immediately after an employee leaves? Does the user question involve sensitive information? Is the cited version outdated? These are not purely model-capability problems; they are system-engineering problems.",
            "Therefore, engineering implementation is not about stacking tools. It is about separating RAG into observable, replaceable, and evaluable layers. Only then can teams locate failures in a specific layer instead of guessing between model, retrieval, and data.",
        ],
    },
    "15": {
        "cn": [
            "六层架构可以作为排障框架。比如答案引用了错误文档，先看数据准备层是否把文档解析错了；再看索引层是否把相关段落切散了；再看检索层是否排序错误；最后看生成层是否误读证据。",
            "每一层都有自己的质量指标。数据层看解析完整率和元数据准确性；索引层看 chunk 质量和增量更新；检索层看召回率、精确率和延迟；生成层看忠实性和引用准确；应用层看权限和交互体验；评估层看回归测试和线上监控。",
            "这页可以强调一个观点：RAG 系统越复杂，越需要分层责任。否则所有问题都会被说成“模型不行”或“检索不行”，团队很难持续优化。",
        ],
        "en": [
            "The six-layer architecture is also a debugging framework. If the answer cites the wrong document, first check whether the data preparation layer parsed the document incorrectly. Then check whether the indexing layer split the relevant passage poorly. Then check whether the retrieval layer ranked results incorrectly. Finally, check whether the generation layer misread the evidence.",
            "Each layer has its own quality metrics. The data layer measures parsing completeness and metadata accuracy. The indexing layer measures chunk quality and incremental updates. The retrieval layer measures recall, precision, and latency. The generation layer measures faithfulness and citation accuracy. The application layer handles permissions and user experience. The evaluation layer handles regression testing and online monitoring.",
            "The main point is that the more complex the RAG system becomes, the more important layered responsibility is. Otherwise, every problem becomes vaguely described as “the model is bad” or “retrieval is bad,” and continuous improvement becomes difficult.",
        ],
    },
    "16": {
        "cn": [
            "工具生态这一页不要讲成工具清单，而要讲成分层选型。比如如果问题是数据接入和索引组织，LlamaIndex 可能更贴近；如果问题是多步骤应用编排和工具调用，LangChain 可能更适合；如果问题是搜索基础设施，就要看向量库和全文检索系统。",
            "评估工具也要尽早引入。很多团队先做系统，后补评估，结果上线后才发现无法判断改动是否变好。RAGAS、LangSmith、DeepEval 这类工具的价值，是把回答质量、上下文质量、链路追踪和回归测试变成可观察对象。",
            "选型时建议从问题倒推工具：数据在哪里、权限怎么管、检索延迟要求多少、是否需要私有化部署、是否要多语言、多模态、Agent 工具调用。工具只是实现手段，架构目标要先明确。",
        ],
        "en": [
            "This tool landscape slide should not be presented as a list of products. It should be presented as layer-based selection. If the issue is data connection and index organization, LlamaIndex may fit well. If the issue is multi-step orchestration and tool calling, LangChain may be more relevant. If the issue is search infrastructure, vector databases and full-text search systems matter more.",
            "Evaluation tools should be introduced early. Many teams build the system first and add evaluation later, only to discover after launch that they cannot tell whether a change improved quality. Tools such as RAGAS, LangSmith, and DeepEval turn answer quality, context quality, tracing, and regression testing into observable objects.",
            "Tool selection should be driven backward from the problem: where the data lives, how permissions are managed, what latency is required, whether private deployment is needed, whether multilingual or multimodal retrieval is needed, and whether Agents must call tools. Tools are implementation choices; the architecture goal should come first.",
        ],
    },
    "17": {
        "cn": [
            "技术范式部分可以告诉听众：RAG 不是一个固定标准答案，而是一组不断演化的系统设计模式。随着任务变复杂，简单的检索生成链路会暴露出召回不稳定、关系理解弱、无法自我修正等问题。",
            "Advanced RAG 主要提升上下文质量；Graph RAG 主要补强关系和全局分析；Modular RAG 让组件可替换；Agentic RAG 让系统能够根据任务动态决策。它们解决的问题不同，所以经常组合使用。",
            "这页的过渡作用很强：前面讲的是经典链路，后面讲的是能力扩展。可以说，理解范式演进能帮助我们判断“什么时候应该增加复杂度”，而不是一开始就把系统做得过重。",
        ],
        "en": [
            "In the technical paradigms section, tell the audience that RAG is not one fixed standard answer. It is a set of evolving system design patterns. As tasks become more complex, a simple retrieve-then-generate chain exposes problems such as unstable recall, weak relationship understanding, and limited self-correction.",
            "Advanced RAG mainly improves context quality. Graph RAG strengthens relationship and global analysis. Modular RAG makes components replaceable. Agentic RAG allows the system to make dynamic decisions based on the task. They solve different problems, so they are often combined.",
            "This slide is a transition: earlier slides covered the classic pipeline, while later slides cover capability extensions. Understanding paradigm evolution helps us decide when complexity is justified, instead of making the system too heavy from the beginning.",
        ],
    },
    "18": {
        "cn": [
            "演进路线可以讲成能力逐步叠加。传统 IR 解决的是“按词找资料”；Dense Retrieval 解决的是“按语义找资料”；Classic RAG 解决的是“找到资料后让模型使用”；Advanced RAG 解决的是“让资料更准、更短、更有用”。",
            "Graph RAG 的出现，是因为很多问题不是单个片段能回答的。例如“某个产品线过去三年有哪些主要风险变化”，需要跨文档关系、时间线和实体关系。纯向量相似度很难直接总结全局结构。",
            "Agentic RAG 则进一步把决策权从固定流程转移到控制器。系统可以根据问题选择是否检索、是否调用工具、是否分解问题。这里要提醒：能力增强的同时，评估、治理和成本控制也必须同步增强。",
        ],
        "en": [
            "The evolution path can be explained as accumulating capabilities. Traditional IR solves “find material by words.” Dense retrieval solves “find material by meaning.” Classic RAG solves “let the model use retrieved material.” Advanced RAG solves “make the material more accurate, shorter, and more useful.”",
            "Graph RAG emerged because many questions cannot be answered by a single passage. For example, “What were the major risk changes for this product line over the past three years?” requires cross-document relationships, timelines, and entity connections. Pure vector similarity does not naturally summarize global structure.",
            "Agentic RAG then shifts decision-making from a fixed pipeline to a controller. The system can decide whether to retrieve, whether to call tools, and whether to decompose the question. The important reminder is that as capability increases, evaluation, governance, and cost control must increase as well.",
        ],
    },
    "19": {
        "cn": [
            "范式地图可以用来避免概念混淆。Naive、Advanced、Corrective、Self、Adaptive、Graph、Agentic 等名字看起来很多，但本质上是在回答不同问题：检索失败怎么办？关系问题怎么办？复杂任务怎么拆？系统如何自检？",
            "生产系统往往不是纯粹某一种范式。例如一个企业客服系统可能用 Hybrid Search 和 Rerank 做 Advanced RAG，用知识图谱处理产品关系，用 Agent 判断是否需要查订单 API，再用评估器检查答案是否忠于证据。",
            "所以这页的结论是：范式不是标签，而是能力模块。真正重要的是识别当前瓶颈，再选择最小必要复杂度来解决它。",
        ],
        "en": [
            "The paradigm map helps avoid conceptual confusion. Names such as Naive, Advanced, Corrective, Self, Adaptive, Graph, and Agentic RAG may look like many separate categories, but they answer different questions: What if retrieval fails? What if relationships matter? How should complex tasks be decomposed? How does the system check itself?",
            "Production systems are rarely pure examples of one paradigm. An enterprise customer-service system may use hybrid search and reranking as Advanced RAG, a knowledge graph for product relationships, an Agent to decide whether to query an order API, and an evaluator to check whether the answer is faithful to evidence.",
            "The conclusion is that paradigms are not labels; they are capability modules. The real task is to identify the current bottleneck and introduce the minimum necessary complexity to solve it.",
        ],
    },
    "20": {
        "cn": [
            "Graph RAG 适合关系密集型问题，比如实体之间的依赖、组织关系、供应链关系、研究主题演化等。它通过抽取实体、关系和社区结构，让系统不只看到相似文本，还能看到知识网络。",
            "Modular RAG 则更偏工程架构。把查询改写、检索、过滤、重排、压缩、生成和评估拆开后，每个模块都可以单独替换和评估。例如 Embedding 模型升级时，不一定要重写整个应用。",
            "两者的共同代价是系统复杂度。图谱需要构建和维护，模块化需要接口约定和监控指标。因此要根据问题价值决定是否引入，而不是因为技术新就引入。",
        ],
        "en": [
            "Graph RAG is suitable for relationship-heavy questions, such as dependencies between entities, organizational relationships, supply-chain relationships, or the evolution of research topics. By extracting entities, relations, and community structures, the system sees not only similar text but also a knowledge network.",
            "Modular RAG is more about engineering architecture. Once query rewriting, retrieval, filtering, reranking, compression, generation, and evaluation are separated, each module can be replaced and evaluated independently. For example, upgrading the embedding model does not necessarily require rewriting the whole application.",
            "The shared cost is complexity. Graphs must be constructed and maintained. Modular systems require interface contracts and monitoring metrics. These approaches should be introduced based on problem value, not simply because the technology is new.",
        ],
    },
    "21": {
        "cn": [
            "Agentic RAG 的关键是让系统具备任务级判断能力。普通 RAG 像固定流水线，来了问题就走同一套流程；Agentic RAG 更像一个调度器，先理解任务，再决定用哪些信息源和工具。",
            "例如简单定义类问题，可能只需要一次知识库检索；订单售后问题，需要先查订单系统，再查政策；复杂分析问题，可能要分解成多个子问题，分别检索、汇总、验证。",
            "但 Agentic 并不意味着无限自由。真正可用的 Agentic RAG 必须有明确的工具边界、权限边界、重试上限和兜底机制。否则系统越聪明，风险也越难控制。",
        ],
        "en": [
            "The key of Agentic RAG is task-level judgment. Classic RAG behaves like a fixed pipeline: every question follows the same route. Agentic RAG behaves more like a dispatcher: it understands the task first, then decides which information sources and tools to use.",
            "For a simple definition question, one knowledge-base retrieval may be enough. For an order after-sales question, the system may need to query the order system first and then retrieve policy documents. For a complex analysis task, it may need to decompose the question, retrieve multiple pieces of evidence, summarize, and verify.",
            "However, Agentic does not mean unlimited freedom. A usable Agentic RAG system must define tool boundaries, permission boundaries, retry limits, and fallback mechanisms. Otherwise, the smarter the system becomes, the harder it is to control risk.",
        ],
    },
    "22": {
        "cn": [
            "这一页可以用“固定路线”和“动态导航”类比。普通 RAG 像固定路线公交车，所有问题都经过同样站点；Agentic RAG 像导航系统，会根据目的地、路况和约束选择路径。",
            "动态能力带来明显好处：可以避免不必要检索，可以在第一次检索失败后改写查询，也可以在知识库不够时调用工具或转人工。但这同时意味着决策本身会成为新的错误来源。",
            "因此，Agentic RAG 的评估对象不只是最终答案，还包括中间决策是否合理：该不该检索？该不该调用工具？工具参数是否正确？是否在证据不足时及时停止？",
        ],
        "en": [
            "This slide can use the analogy of a fixed route versus dynamic navigation. Classic RAG is like a bus route: every question goes through the same stops. Agentic RAG is like a navigation system: it chooses a path based on destination, conditions, and constraints.",
            "Dynamic capability brings clear benefits. It can avoid unnecessary retrieval, rewrite the query after a failed retrieval, call tools when the knowledge base is insufficient, or hand off to a human. But this also means that decision-making itself becomes a new source of errors.",
            "Therefore, evaluating Agentic RAG is not only about the final answer. We also need to evaluate intermediate decisions: Should the system retrieve? Should it call a tool? Are tool parameters correct? Does it stop when evidence is insufficient?",
        ],
    },
    "23": {
        "cn": [
            "Agent Loop 可以讲成观察、判断、行动的闭环。系统先理解问题和约束，再选择一个动作，比如检索文档、查询 API、调用计算工具或请求澄清。得到结果后，它评估证据质量，再决定下一步。",
            "证据评估是闭环的关键。如果检索结果只是表面相关，系统应该重写查询；如果多个来源冲突，系统应该提示冲突；如果证据足够，就应该停止检索并生成答案，而不是无限追加上下文。",
            "为了让闭环可控，必须加工程护栏：最大轮数、工具白名单、参数校验、权限检查、日志追踪和人工兜底。没有这些护栏，Agentic RAG 很容易从增强能力变成不可预测行为。",
        ],
        "en": [
            "The Agent loop can be explained as a cycle of observe, judge, and act. The system first understands the question and constraints, then chooses an action, such as retrieving documents, querying an API, calling a calculation tool, or asking for clarification. After receiving results, it evaluates evidence quality and decides the next step.",
            "Evidence evaluation is the key part of the loop. If retrieval results are only superficially relevant, the system should rewrite the query. If sources conflict, it should report the conflict. If evidence is sufficient, it should stop retrieving and generate the answer instead of adding endless context.",
            "To keep the loop controllable, engineering guardrails are required: maximum rounds, tool allowlists, parameter validation, permission checks, logging, and human fallback. Without these guardrails, Agentic RAG can turn from an enhancement into unpredictable behavior.",
        ],
    },
    "24": {
        "cn": [
            "角色分工本质上是把 Agent 的复杂行为拆成可控模块。路由器决定问题走哪条路；规划器决定任务怎么拆；检索控制器决定查哪里；工具调用者负责和外部系统交互；证据评估器负责判断材料质量。",
            "生成控制器和反思器也很重要。生成控制器确保模型按证据回答，并输出合适格式；反思器检查答案是否遗漏、是否引用错误、是否超出证据。兜底处理器则保证系统在不确定时不会硬答。",
            "这页可以强调：Agentic RAG 不是让一个大模型随意做所有事，而是把能力拆成角色，再用边界和监控管理每个角色。可控性来自清晰分工。",
        ],
        "en": [
            "Role separation turns complex Agent behavior into controllable modules. The router decides which path the question should take. The planner decides how to decompose the task. The retrieval controller decides where to search. The tool caller interacts with external systems. The evidence evaluator judges evidence quality.",
            "The generation controller and reflector are also important. The generation controller ensures the model answers from evidence and uses the right format. The reflector checks whether the answer misses key points, cites incorrectly, or goes beyond the evidence. The fallback handler prevents the system from forcing an answer when uncertainty is high.",
            "The key message is that Agentic RAG is not about letting one model do everything freely. It is about separating capabilities into roles, then managing each role with boundaries and monitoring. Controllability comes from clear responsibility.",
        ],
    },
    "25": {
        "cn": [
            "这个案例可以讲得更像真实客服流程。用户的问题里包含两个信息：订单还在运输中，以及商品坏了。系统首先要识别订单号 A1001，然后调用订单系统确认状态，而不是仅凭用户描述判断。",
            "接着系统要识别售后意图，检索退款政策和故障处理政策。政策可能规定“未签收不能申请某类退款”，也可能规定“质量问题需提供照片或检测结果”。这些条件需要和订单状态组合判断。",
            "最后生成答案时，系统不能简单说“可以”或“不可以”，而要说明目前依据是什么、还缺什么信息、下一步该怎么做。如果证据不足，就请求补充或转人工。这体现了 Agentic RAG 的核心价值：在复杂业务中做有边界的决策。",
        ],
        "en": [
            "This case can be presented like a real customer-service workflow. The user question contains two pieces of information: the order is still in transit, and the headphones are broken. The system first identifies order ID A1001 and calls the order system to confirm the status, instead of trusting the user description alone.",
            "Next, the system identifies the after-sales intent and retrieves refund and fault-handling policies. The policy may say that certain refunds are unavailable before delivery, or that quality issues require photos or inspection results. These conditions must be combined with order status.",
            "When generating the final answer, the system should not simply say yes or no. It should explain what evidence is available, what information is missing, and what the next step should be. If evidence is insufficient, it should ask for clarification or hand off to a human. This shows the core value of Agentic RAG: bounded decision-making in complex business workflows.",
        ],
    },
    "26": {
        "cn": [
            "Agentic RAG 的优势来自更强的任务适应性。它能判断简单问题直接回答、知识问题检索、业务问题查工具、风险问题转人工，也能在检索失败后主动修正路线。",
            "但它的风险同样来自动态性。每多一轮检索、多一次工具调用、多一次模型判断，就多一个成本点和错误点。尤其工具调用涉及真实业务系统时，错误参数可能带来真实后果。",
            "上线时建议设置几类指标：平均轮数、工具调用成功率、无效调用率、超时率、人工转接率、答案忠实性和用户反馈。Agentic RAG 不能只按回答好不好看评估，还要按流程是否可靠评估。",
        ],
        "en": [
            "The advantage of Agentic RAG comes from stronger task adaptability. It can answer simple questions directly, retrieve for knowledge questions, call tools for business questions, escalate risky questions to humans, and adjust the route after retrieval failure.",
            "But the risks also come from this dynamic behavior. Every additional retrieval round, tool call, and model judgment adds cost and another possible failure point. When tool calls involve real business systems, wrong parameters may create real consequences.",
            "For production, teams should monitor metrics such as average number of rounds, tool-call success rate, invalid-call rate, timeout rate, human handoff rate, answer faithfulness, and user feedback. Agentic RAG should not be evaluated only by whether the final answer looks good; the process itself must be reliable.",
        ],
    },
    "27": {
        "cn": [
            "应用、挑战和评估这一部分可以作为落地总结。RAG 的价值只有在具体场景中才成立，因为不同场景对准确性、时效性、权限和引用的要求不同。",
            "比如企业知识库强调版本和权限，客服强调流程和承诺边界，法律医疗强调引用和人工复核，金融强调实时性和合规，代码助手强调项目结构和依赖关系。没有场景约束，就很难定义什么叫“好答案”。",
            "评估也要跟场景绑定。低风险 FAQ 可以自动评估为主；高风险法律医疗需要专家抽检；生产客服还要看用户满意度、转人工率、响应时间和错误承诺率。",
        ],
        "en": [
            "The applications, challenges, and evaluation section can serve as the implementation summary. RAG’s value only becomes meaningful in concrete scenarios, because different scenarios have different requirements for accuracy, freshness, permissions, and citations.",
            "For example, enterprise knowledge bases emphasize versions and permissions. Customer service emphasizes workflows and commitment boundaries. Law and medicine emphasize citations and expert review. Finance emphasizes timeliness and compliance. Code assistants emphasize project structure and dependencies. Without scenario constraints, it is hard to define a “good answer.”",
            "Evaluation must also be scenario-specific. Low-risk FAQs may rely mainly on automated evaluation. High-risk legal and medical use cases require expert sampling. Production customer service should also measure user satisfaction, handoff rate, response time, and incorrect commitment rate.",
        ],
    },
    "28": {
        "cn": [
            "典型场景这一页可以按“为什么适合”和“风险边界”两条线讲。企业知识库适合 RAG，因为文档频繁更新且权限复杂；但必须保证用户只能检索自己有权看的内容。",
            "客服场景适合 RAG，因为 FAQ、政策和订单系统可以组合；但不能让模型自行承诺退款、赔偿或时效。法律医疗适合 RAG，因为答案必须有依据；但也最需要人工复核，不能把模型当最终决策者。",
            "代码助手和多模态检索说明 RAG 不只处理普通文本。代码需要理解仓库结构和版本，图纸表格截图需要 OCR、布局解析和多模态 embedding。RAG 的范围正在从文档问答扩展到企业知识操作系统。",
        ],
        "en": [
            "For the use-case slide, explain each scenario through two lines: why RAG fits, and where the risk boundary is. Enterprise knowledge bases fit RAG because documents update frequently and permissions are complex. But the system must ensure users only retrieve content they are allowed to see.",
            "Customer service fits RAG because FAQs, policies, and order systems can be combined. But the model must not independently promise refunds, compensation, or delivery times. Law and medicine fit RAG because answers require evidence, but they also require the strongest human review. The model should not be treated as the final decision-maker.",
            "Code assistants and multimodal retrieval show that RAG is not limited to ordinary text. Code requires repository structure and version understanding. Drawings, tables, and screenshots require OCR, layout parsing, and multimodal embeddings. RAG is expanding from document QA toward an enterprise knowledge operating system.",
        ],
    },
    "29": {
        "cn": [
            "优势这一页可以总结成四个关键词：更新、私有、可追溯、可控。更新意味着知识变化时更新文档和索引即可；私有意味着企业内部知识不必进入模型训练；可追溯意味着答案能返回来源；可控意味着权限和过滤可以工程化实现。",
            "降低幻觉不是因为模型突然不会犯错，而是因为模型被要求在证据范围内回答。换句话说，RAG 通过缩小模型自由发挥的空间，换取更高的事实可靠性。",
            "也可以提醒：RAG 不能保证 100% 正确。它降低的是某些错误概率，并提供排查机制。如果证据库本身错了、检索漏了、引用错了，RAG 仍会失败，所以优势必须和评估、监控一起看。",
        ],
        "en": [
            "The advantages can be summarized with four keywords: updateability, privacy, traceability, and controllability. Updateability means knowledge changes can be handled by updating documents and indexes. Privacy means internal knowledge does not need to become training data. Traceability means answers can return sources. Controllability means permissions and filters can be engineered.",
            "RAG reduces hallucination not because the model suddenly becomes incapable of mistakes, but because the model is constrained to answer within the evidence. In other words, RAG trades some freedom for higher factual reliability.",
            "It is also important to remind the audience that RAG does not guarantee 100 percent correctness. It reduces certain error probabilities and provides a way to investigate failures. If the evidence base is wrong, retrieval misses the right document, or citations are wrong, RAG can still fail. Its advantages must be considered together with evaluation and monitoring.",
        ],
    },
    "30": {
        "cn": [
            "挑战这一页适合强调“不要把问题都归因给模型”。如果正确文档没有被召回，应该改检索；如果召回内容太杂，应该改重排和压缩；如果答案没有引用来源，应该改 Prompt 和输出格式；如果越权访问，应该改权限和数据治理。",
            "很多挑战是互相牵连的。Chunk 太小可能提高精确度但损失上下文；Top-K 太大可能提高召回但增加噪声；压缩可以降低成本但可能丢掉关键细节。因此优化 RAG 常常是平衡问题。",
            "建议讲一个排查顺序：先建立标准测试集，然后检查入库、召回、排序、上下文、生成、引用和权限。每次只改一个主要变量，用回归评估确认是否真的变好。",
        ],
        "en": [
            "This challenge slide should emphasize that not every problem should be blamed on the model. If the correct document is not retrieved, improve retrieval. If retrieved content is noisy, improve reranking and compression. If the answer lacks citations, improve the prompt and output format. If unauthorized access occurs, fix permissions and data governance.",
            "Many challenges interact with each other. Smaller chunks may improve precision but lose context. Larger Top-K may improve recall but add noise. Compression can reduce cost but may remove key details. Optimizing RAG is often a balancing problem.",
            "A practical debugging order is: build a standard test set, then check indexing, recall, ranking, context assembly, generation, citation, and permissions. Change one major variable at a time and use regression evaluation to confirm whether quality actually improves.",
        ],
    },
    "31": {
        "cn": [
            "评估这一页要讲清楚“答案对了”并不等于系统没问题。答案可能碰巧正确，但引用错了；引用可能正确，但检索漏掉了更关键文档；检索可能正确，但模型没有忠于上下文。",
            "因此检索指标和生成指标必须分开。检索侧关注正确证据是否进入 Top-K、相关内容召回是否充分、噪声是否太多、排序是否合理。生成侧关注答案是否忠实、是否回答问题、是否事实正确、引用是否支持结论。",
            "线上监控也不能只看自动分数。需要结合用户反馈、人工抽检、成本、延迟、拒答率、转人工率和失败案例分析。RAG 评估最好形成持续回归机制，每次改索引、模型或 Prompt 后都能比较变化。",
        ],
        "en": [
            "This evaluation slide should make clear that a correct-looking answer does not mean the system is healthy. The answer may be accidentally correct but cite the wrong source. The citation may be correct but retrieval may have missed a more important document. Retrieval may be correct but the model may not remain faithful to the context.",
            "Therefore, retrieval metrics and generation metrics must be separated. Retrieval-side metrics ask whether correct evidence enters Top-K, whether enough relevant context is recalled, whether there is too much noise, and whether ranking is reasonable. Generation-side metrics ask whether the answer is faithful, relevant, factually correct, and properly supported by citations.",
            "Online monitoring should not rely only on automated scores. It should combine user feedback, human sampling, cost, latency, refusal rate, handoff rate, and failure-case analysis. Ideally, RAG evaluation becomes a continuous regression mechanism, so every change to the index, model, or prompt can be compared.",
        ],
    },
    "32": {
        "cn": [
            "这一页可以帮助听众避免技术路线混用。Fine-tuning 适合改模型行为，比如语气、格式、领域表达和任务习惯；但如果知识每天变，靠微调更新事实成本很高。",
            "长上下文可以放更多资料，但“能放下”不等于“能用好”。没有检索和组织，长上下文会带来噪声、成本和注意力分散。Prompt Engineering 决定模型如何使用证据，但 Prompt 不能替代证据质量。",
            "知识图谱适合结构化关系，Agent 适合规划和工具调用。RAG 可以成为 Agent 的一个工具，也可以结合图谱增强关系理解。关键是让每种技术承担合适职责，而不是让单一技术解决所有问题。",
        ],
        "en": [
            "This slide helps the audience avoid mixing up technical routes. Fine-tuning is suitable for changing model behavior, such as tone, format, domain expression, and task habits. But if facts change every day, using fine-tuning to update knowledge is expensive.",
            "Long context can hold more material, but being able to fit more text does not mean the model can use it well. Without retrieval and organization, long context brings noise, cost, and attention dispersion. Prompt engineering defines how the model uses evidence, but prompts cannot replace evidence quality.",
            "Knowledge graphs are suitable for structured relationships. Agents are suitable for planning and tool use. RAG can become one tool used by an Agent, and it can also combine with graphs to improve relationship understanding. The key is to assign the right responsibility to each technique instead of forcing one technique to solve everything.",
        ],
    },
    "33": {
        "cn": [
            "最后一页可以把未来趋势和前面的主线连接起来。未来 RAG 不是简单多接几个文档源，而是变得更动态、更结构化、更实时、更受权限约束，也更依赖自动评估。",
            "Agentic RAG 会增强动态决策，Graph RAG 会增强关系理解，多模态 RAG 会扩展知识形态，实时 RAG 会连接业务系统和数据流，个性化与权限感知 RAG 会让检索结果随身份和场景变化。",
            "结束时可以用一句话收束：RAG 的终极目标不是让模型“知道更多”，而是让模型在正确的时间、基于正确权限、使用正确证据，给出可解释、可追溯、可控制的答案。",
        ],
        "en": [
            "The final slide should connect future trends back to the main storyline. Future RAG is not simply about connecting more document sources. It will become more dynamic, more structured, more real-time, more permission-aware, and more dependent on automated evaluation.",
            "Agentic RAG will strengthen dynamic decision-making. Graph RAG will strengthen relationship understanding. Multimodal RAG will expand the forms of knowledge. Real-time RAG will connect business systems and data streams. Personalized and permission-aware RAG will make retrieval results depend on identity and context.",
            "A strong closing sentence is: the ultimate goal of RAG is not to make the model “know more,” but to make the model use the right evidence, under the right permissions, at the right time, and produce answers that are explainable, traceable, and controllable.",
        ],
    },
}


def set_font(run, size=None, bold=False, color=None, east_asia="STHeiti"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_para(doc, text, size=10.2, color=None, bold=False, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    color = "2E74B5" if level == 1 else "1F4D78"
    size = 15 if level == 1 else 12
    for run in p.runs:
        set_font(run, size=size, bold=True, color=color)
    return p


def paragraphs_for(slide, lang):
    items = [slide[lang]]
    items.extend(details[slide["no"]][lang])
    return items


def build_md():
    lines = [
        "# RAG 中英双语展示讲稿（详细版）",
        "",
        "说明：本稿按英文版 PPT 的 33 页逐页对应。每页包含中文详细讲稿与英文详细讲稿，并在原讲稿基础上补充解释、例子、风险点和讲法提示。",
        "",
    ]
    for s in slides:
        lines.extend([
            f"## Slide {s['no']} - {s['title']}",
            "",
            "### 中文详细讲稿",
            "",
            *paragraphs_for(s, "cn"),
            "",
            "### English Detailed Script",
            "",
            *paragraphs_for(s, "en"),
            "",
        ])
    return "\n\n".join(lines)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.22

    header = section.header.paragraphs[0]
    header.text = "RAG bilingual detailed speaker script"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=8.5, color="666666")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("RAG 中英双语展示讲稿（详细版）")
    set_font(r, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("逐页对应英文版 PPT，共 33 页。每页包含中文详细讲稿与英文详细讲稿，补充解释、例子、风险点和转场。")
    set_font(r, size=10.5, color="666666")

    meta = add_para(doc, "对应文件：RAG_Technical_Review_EN.pptx；内容原则：保持页序和主题不变，在每页内扩展可讲细节。", size=9.5, color="666666", after=10)
    shade_paragraph(meta, "F4F6F9")

    for s in slides:
        add_heading(doc, f"Slide {s['no']}  {s['title']}", 1)
        add_heading(doc, "中文详细讲稿", 2)
        for para in paragraphs_for(s, "cn"):
            add_para(doc, para)
        add_heading(doc, "English Detailed Script", 2)
        for para in paragraphs_for(s, "en"):
            add_para(doc, para)

    doc.save(DOCX)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    MD.write_text(build_md(), encoding="utf-8")
    build_docx()
    print(DOCX)
    print(MD)


if __name__ == "__main__":
    main()
