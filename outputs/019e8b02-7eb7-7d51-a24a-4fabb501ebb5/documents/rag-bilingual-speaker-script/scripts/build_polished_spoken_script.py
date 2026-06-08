from pathlib import Path
import importlib.util

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SPOKEN_SCRIPT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/scripts/build_spoken_detailed_script.py")
OUT = Path("/Users/caiwenbin/Desktop/项目/agentic rag/outputs/019e8b02-7eb7-7d51-a24a-4fabb501ebb5/documents/rag-bilingual-speaker-script/output")
DOCX = OUT / "RAG_中英双语展示讲稿_自然过渡润色版.docx"
MD = OUT / "RAG_中英双语展示讲稿_自然过渡润色版.md"


spec = importlib.util.spec_from_file_location("spoken", SPOKEN_SCRIPT)
spoken = importlib.util.module_from_spec(spec)
spec.loader.exec_module(spoken)
slides = spoken.slides


transitions = {
    "01": {
        "cn": "各位好，今天我们从一个很实际的问题开始：大模型已经很强了，为什么还需要 RAG？带着这个问题，我们先看整场汇报的主题和结构。",
        "en": "Hello everyone. Let us begin with a practical question: large language models are already powerful, so why do we still need RAG? With that question in mind, let us start with the theme and structure of the talk.",
    },
    "02": {
        "cn": "刚才我们看到 RAG 的整体定位，接下来先不要急着进入技术细节，而是先把它为什么出现讲清楚。",
        "en": "Now that we have positioned RAG at a high level, let us not rush into implementation yet. First, we need to understand why RAG exists.",
    },
    "03": {
        "cn": "有了动机之后，我们就可以给 RAG 一个更准确的定义：它到底在系统里做了什么，数据又是怎样流动的。",
        "en": "With the motivation in place, we can now define RAG more precisely: what it does inside the system and how information flows through it.",
    },
    "04": {
        "cn": "理解了定义之后，我们再往下问一步：如果没有 RAG，只让模型依赖参数记忆，会出现哪些具体问题？",
        "en": "After defining RAG, the next question is: what exactly goes wrong if we do not use RAG and rely only on the model’s parameter memory?",
    },
    "05": {
        "cn": "这些问题背后，其实对应着两类知识的分工。讲清楚这点，后面的工程链路就更容易理解。",
        "en": "Behind these problems is a division of labor between two types of knowledge. Once this is clear, the engineering pipeline becomes much easier to understand.",
    },
    "06": {
        "cn": "到这里，概念部分就收住了。接下来我们进入第二部分，看 RAG 在工程上到底由哪些链路组成。",
        "en": "That closes the conceptual foundation. Next, we move into the second section and look at the engineering pipeline behind RAG.",
    },
    "07": {
        "cn": "先看全链路。只有把离线和在线两部分放在一起看，才能判断一个 RAG 系统为什么好，或者为什么会失败。",
        "en": "Let us first look at the full pipeline. We need to see offline and online stages together before we can understand why a RAG system works or fails.",
    },
    "08": {
        "cn": "在这条链路里，第一块基础工作是离线索引。它决定了知识能不能被后续系统真正找到。",
        "en": "Within this pipeline, the first foundational stage is offline indexing. It determines whether knowledge can actually be found later.",
    },
    "09": {
        "cn": "离线索引里最容易被低估的一步，是文档切分。它看起来像预处理，实际上会直接影响检索和生成。",
        "en": "One of the most underestimated steps in offline indexing is chunking. It may look like preprocessing, but it directly affects both retrieval and generation.",
    },
    "10": {
        "cn": "切分解决的是知识颗粒度，Embedding 解决的则是语义匹配。接下来我们看系统怎样理解“相似”和“相关”。",
        "en": "Chunking defines the granularity of knowledge. Embedding defines semantic matching. Now let us look at how the system understands similarity and relevance.",
    },
    "11": {
        "cn": "当文档被向量化之后，下一个问题就是规模。数据一多，如何又快又准地查到相关向量，就需要向量数据库来解决。",
        "en": "Once documents are embedded, the next issue is scale. When the dataset grows, we need vector databases to find relevant vectors quickly and accurately.",
    },
    "12": {
        "cn": "有了向量库之后，我们再回到检索本身。实际系统很少只靠一种检索方式，原因就在于不同方法会漏掉不同类型的问题。",
        "en": "After discussing vector databases, let us return to retrieval itself. Real systems rarely rely on only one retrieval method, because different methods fail in different ways.",
    },
    "13": {
        "cn": "检索到候选内容还不够，真正进入模型上下文之前，还需要一次筛选、整理和约束，这就是重排序和 Prompt 组织的作用。",
        "en": "Retrieving candidate content is not enough. Before evidence enters the model context, it needs to be selected, organized, and constrained. That is where reranking and prompt assembly come in.",
    },
    "14": {
        "cn": "到这里，核心链路已经讲完。下面我们把视角从单条链路扩大到生产系统，看一个可用的 RAG 应该怎样工程化。",
        "en": "At this point, we have covered the core pipeline. Now let us zoom out from a single chain to a production system and see how RAG should be engineered.",
    },
    "15": {
        "cn": "工程化的第一步，是把系统分层。分层不是为了画架构图好看，而是为了让责任、指标和排障路径更清楚。",
        "en": "The first step in engineering is layering the system. This is not just for a clean architecture diagram; it clarifies responsibilities, metrics, and debugging paths.",
    },
    "16": {
        "cn": "分层之后，工具选型就自然多了。我们不再问哪个工具最好，而是问当前缺的是哪一层能力。",
        "en": "Once the layers are clear, tool selection becomes more natural. We no longer ask which tool is best; we ask which layer capability is missing.",
    },
    "17": {
        "cn": "工程实现讲完后，我们再看技术范式的演进。因为 RAG 不是一个静态方案，而是在不断解决上一代系统的问题。",
        "en": "After engineering implementation, let us look at the evolution of RAG paradigms. RAG is not a static solution; it keeps evolving to fix the weaknesses of earlier systems.",
    },
    "18": {
        "cn": "先把演进脉络拉直：每一代 RAG 并不是推翻上一代，而是在上一代能力上补短板。",
        "en": "Let us first clarify the evolution path. Each generation of RAG does not completely replace the previous one; it fills gaps left by the previous generation.",
    },
    "19": {
        "cn": "明白演进之后，我们用这张图把不同范式放到一张地图里，避免把它们理解成互相排斥的选项。",
        "en": "With that evolution in mind, this map places different paradigms together so we do not misunderstand them as mutually exclusive choices.",
    },
    "20": {
        "cn": "在这些范式里，Graph RAG 和 Modular RAG 很有代表性。它们解决的不是普通相似度检索，而是关系理解和工程组合的问题。",
        "en": "Among these paradigms, Graph RAG and Modular RAG are especially representative. They address relationship understanding and engineering composition, not just similarity retrieval.",
    },
    "21": {
        "cn": "讲到这里，我们自然过渡到 Agentic RAG。它进一步回答一个问题：当任务复杂起来时，谁来决定下一步该怎么做？",
        "en": "This naturally leads us to Agentic RAG. It answers the next question: when the task becomes complex, who decides what the system should do next?",
    },
    "22": {
        "cn": "为了理解 Agentic RAG，我们先把它和普通 RAG 放在一起对比。差异不是多一个组件，而是决策方式变了。",
        "en": "To understand Agentic RAG, let us compare it with classic RAG. The difference is not just one additional component; the decision-making mechanism changes.",
    },
    "23": {
        "cn": "既然 Agentic RAG 的核心是决策，那么接下来就要看它怎样在观察、判断和行动之间循环。",
        "en": "Since the core of Agentic RAG is decision-making, the next step is to look at how it loops through observation, judgment, and action.",
    },
    "24": {
        "cn": "有了循环之后，还需要把循环里的职责拆开。否则一个 Agent 什么都做，就很难控制和排查。",
        "en": "Once we have the loop, we also need to separate responsibilities inside it. If one Agent does everything, control and debugging become difficult.",
    },
    "25": {
        "cn": "这些角色听起来比较抽象，所以我们用一个客服问题把它们串起来，看 Agentic RAG 为什么有价值。",
        "en": "These roles may sound abstract, so let us connect them through a customer-service case and see why Agentic RAG is valuable.",
    },
    "26": {
        "cn": "案例展示了能力，接下来也要看代价。Agentic RAG 更灵活，但灵活性本身也会带来新的风险。",
        "en": "The case shows the capability, but we also need to look at the cost. Agentic RAG is more flexible, and that flexibility introduces new risks.",
    },
    "27": {
        "cn": "Agentic RAG 讲完后，我们进入最后一部分：应用、挑战和评估。也就是把前面的技术落到真实场景里检验。",
        "en": "After Agentic RAG, we move into the final section: applications, challenges, and evaluation. This is where the previous techniques are tested in real scenarios.",
    },
    "28": {
        "cn": "先看应用场景。判断一个场景是否适合 RAG，不是看它是否热门，而是看它是否需要私有、变化快、可引用的知识。",
        "en": "Let us start with use cases. A scenario is suitable for RAG not because it is popular, but because it needs private, fast-changing, and citable knowledge.",
    },
    "29": {
        "cn": "从这些场景里，我们可以抽象出 RAG 的优势。它真正带来的不是一个更大的模型，而是一套管理证据的工程能力。",
        "en": "From these scenarios, we can extract RAG’s advantages. What it really provides is not a larger model, but an engineering capability for managing evidence.",
    },
    "30": {
        "cn": "当然，有优势就会有挑战。RAG 的难点往往不是单点问题，而是链路中多个环节共同作用的结果。",
        "en": "Of course, advantages come with challenges. RAG problems are often not single-point failures, but the result of multiple pipeline stages interacting.",
    },
    "31": {
        "cn": "因此，最后必须讲评估。没有评估，RAG 系统就很难知道自己到底是变好了，还是只是看起来变好了。",
        "en": "That is why evaluation is essential. Without evaluation, a RAG system cannot tell whether it has actually improved or only looks better.",
    },
    "32": {
        "cn": "评估之外，还要把 RAG 和其他技术的关系讲清楚。这样我们才不会把所有问题都强行交给同一种方案。",
        "en": "Beyond evaluation, we also need to clarify how RAG relates to other techniques. This prevents us from forcing every problem into one solution.",
    },
    "33": {
        "cn": "最后，我们用未来趋势收束。前面讲的是 RAG 今天怎样工作，这一页讲的是它接下来会往哪里走。",
        "en": "Finally, let us close with future trends. The previous slides explained how RAG works today; this slide looks at where it is heading next.",
    },
}


def smoothen_cn(text):
    replacements = [
        ("这一页是整场展示的开场。", "先看这一页。"),
        ("这一页进入第一部分：", "从这一页开始，我们进入第一部分："),
        ("RAG 可以概括为：", "把前面的动机收束起来，RAG 可以概括为："),
        ("这一页解释 RAG 解决的五类核心问题。", "具体来看，RAG 主要解决五类核心问题。"),
        ("这里要区分两类知识。", "接下来，我们区分两类知识。"),
        ("第二部分开始讲 RAG 的核心链路。", "进入第二部分后，我们开始看 RAG 的核心链路。"),
        ("这一页把生产级 RAG 的端到端流程串起来：", "把流程串起来看，生产级 RAG 通常包括："),
        ("离线索引更像数据工程，而不是 Prompt 工程。", "先说离线索引。它更像数据工程，而不是 Prompt 工程。"),
        ("Embedding 决定语义检索能否理解领域语言。", "切分之后，就进入 Embedding。Embedding 决定语义检索能否理解领域语言。"),
        ("当数据规模变大时，", "当这些向量规模变大时，"),
        ("这一页对比 BM25、Dense Retrieval 和 Hybrid Search。", "回到检索方法本身，这一页对比 BM25、Dense Retrieval 和 Hybrid Search。"),
        ("第三部分进入工程实现。", "第三部分，我们进入工程实现。"),
        ("一个实用 RAG 系统通常可以分成六层。", "具体拆开看，一个实用 RAG 系统通常可以分成六层。"),
        ("工具选型要按层理解，", "沿着这个分层视角，工具选型也要按层理解，"),
        ("第四部分讲技术范式。", "第四部分，我们看技术范式。"),
        ("RAG 的演进可以看作是在不断修复上一代系统的薄弱点。", "沿着这条线看，RAG 的演进可以理解为不断修复上一代系统的薄弱点。"),
        ("这一页强调，", "所以这一页强调，"),
        ("第五部分进入 Agentic RAG。", "第五部分，我们进入 Agentic RAG。"),
        ("Agentic RAG 的差异不在于多一个 Agent，", "对比来看，Agentic RAG 的差异不在于多一个 Agent，"),
        ("这一页展示 Agentic RAG 的循环。", "接下来这一页展示 Agentic RAG 的循环。"),
        ("Agentic RAG 要灵活又可控，", "要让 Agentic RAG 既灵活又可控，"),
        ("这个客服案例说明", "这个客服案例把前面的角色串了起来，说明"),
        ("Agentic RAG 的优势是", "从能力上看，Agentic RAG 的优势是"),
        ("第六部分把讨论落到应用、挑战和评估。", "第六部分，我们把讨论落到应用、挑战和评估。"),
        ("RAG 最适合", "先看适用场景。RAG 最适合"),
        ("RAG 的优势来自证据外置，", "总结这些场景可以看到，RAG 的优势来自证据外置，"),
        ("RAG 的挑战必须按链路定位。", "讲完优势，也要看到挑战。RAG 的挑战必须按链路定位。"),
        ("RAG 评估必须拆开检索质量和生成质量。", "因此，RAG 评估必须拆开检索质量和生成质量。"),
        ("RAG 不是要替代所有技术，", "最后再看技术分工。RAG 不是要替代所有技术，"),
        ("最后一页总结未来趋势。", "最后一页，我们总结未来趋势。"),
    ]
    out = text
    for old, new in replacements:
        out = out.replace(old, new)
    return out


def smoothen_en(text):
    replacements = [
        ("This opening slide sets the scope", "Let us start with this opening slide, which sets the scope"),
        ("This slide introduces the first section:", "From here, we enter the first section:"),
        ("RAG can be summarized as follows:", "With that motivation in mind, RAG can be summarized as follows:"),
        ("This slide explains five core problems", "More specifically, this slide explains five core problems"),
        ("Here we distinguish two types of knowledge.", "Next, we distinguish two types of knowledge."),
        ("The second section begins with the core RAG pipeline.", "In the second section, we begin with the core RAG pipeline."),
        ("This slide connects the production RAG flow end to end:", "If we connect the production flow end to end, it includes"),
        ("Offline indexing is closer to data engineering", "Let us start with offline indexing. It is closer to data engineering"),
        ("Embedding quality determines", "After chunking, we move to embeddings. Embedding quality determines"),
        ("As data grows,", "As these vectors grow in number,"),
        ("This slide compares BM25", "Returning to retrieval methods, this slide compares BM25"),
        ("The third section moves into engineering implementation.", "The third section moves us into engineering implementation."),
        ("A practical RAG system usually has six layers.", "More concretely, a practical RAG system usually has six layers."),
        ("Tool selection should be understood by layer", "Following this layered view, tool selection should be understood by layer"),
        ("The fourth section discusses technical paradigms.", "In the fourth section, we look at technical paradigms."),
        ("The evolution of RAG can be seen", "Following this line of development, the evolution of RAG can be seen"),
        ("This slide emphasizes that", "So this slide emphasizes that"),
        ("The fifth section introduces Agentic RAG.", "The fifth section brings us to Agentic RAG."),
        ("The difference in Agentic RAG is not", "In comparison, the difference in Agentic RAG is not"),
        ("This slide shows the Agentic RAG loop.", "The next slide shows the Agentic RAG loop."),
        ("For Agentic RAG to be flexible", "To make Agentic RAG both flexible"),
        ("This customer-service case shows", "This customer-service case connects the previous roles and shows"),
        ("Agentic RAG is useful because", "From a capability perspective, Agentic RAG is useful because"),
        ("The sixth section brings the discussion", "In the sixth section, we bring the discussion"),
        ("RAG is strongest when", "Let us start with applicable scenarios. RAG is strongest when"),
        ("RAG’s advantages come from externalizing evidence", "Across these scenarios, we can see that RAG’s advantages come from externalizing evidence"),
        ("RAG challenges must be located", "After the advantages, we also need to look at the challenges. RAG challenges must be located"),
        ("RAG evaluation must separate", "This is why RAG evaluation must separate"),
        ("RAG is not meant to replace", "Finally, let us clarify the technical division of labor. RAG is not meant to replace"),
        ("The final slide summarizes future trends.", "On the final slide, we summarize future trends."),
    ]
    out = text
    for old, new in replacements:
        out = out.replace(old, new)
    return out


def slide_paragraphs(slide, lang):
    items = [transitions[slide["no"]][lang]]
    source = spoken.slide_paragraphs(slide, lang)
    transform = smoothen_cn if lang == "cn" else smoothen_en
    items.extend(transform(p) for p in source)
    return items


def set_font(run, size=None, bold=False, color=None, east_asia="STHeiti"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_para(doc, text, size=10.2, color=None, bold=False, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    color = "2E74B5" if level == 1 else "1F4D78"
    size = 15 if level == 1 else 12
    for run in p.runs:
        set_font(run, size=size, bold=True, color=color)
    return p


def build_md():
    lines = [
        "# RAG 中英双语展示讲稿（自然过渡润色版）",
        "",
        "说明：本稿按英文版 PPT 的 33 页逐页对应。相比口播详细版，本版增加页间承接和段落衔接，使表达更像连续演讲。",
        "",
    ]
    for s in slides:
        lines.extend([
            f"## Slide {s['no']} - {s['title']}",
            "",
            "### 中文口播稿",
            "",
            *slide_paragraphs(s, "cn"),
            "",
            "### English Spoken Script",
            "",
            *slide_paragraphs(s, "en"),
            "",
        ])
    return "\n\n".join(lines)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.22

    header = section.header.paragraphs[0]
    header.text = "RAG bilingual polished spoken script"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=8.5, color="666666")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("RAG 中英双语展示讲稿（自然过渡润色版）")
    set_font(r, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("逐页对应英文版 PPT，共 33 页。增加页间承接和段落衔接，使表达更像连续演讲。")
    set_font(r, size=10.5, color="666666")

    meta = add_para(doc, "对应文件：RAG_Technical_Review_EN.pptx；处理原则：页序和主题不变，重点润色过渡、承接和口播自然度。", size=9.5, color="666666", after=10)
    shade_paragraph(meta, "F4F6F9")

    for s in slides:
        add_heading(doc, f"Slide {s['no']}  {s['title']}", 1)
        add_heading(doc, "中文口播稿", 2)
        for para in slide_paragraphs(s, "cn"):
            add_para(doc, para)
        add_heading(doc, "English Spoken Script", 2)
        for para in slide_paragraphs(s, "en"):
            add_para(doc, para)

    doc.save(DOCX)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    MD.write_text(build_md(), encoding="utf-8")
    build_docx()
    print(DOCX)
    print(MD)


if __name__ == "__main__":
    main()
